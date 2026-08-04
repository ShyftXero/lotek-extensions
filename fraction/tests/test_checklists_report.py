"""Checklists in the report: coverage section + compliance attestation appendix, honoring
include_in_report and cross-linking a failed coverage item to its finding."""

from __future__ import annotations

import io

from fraction import checklists as C
from fraction.models import ChecklistTemplate, Engagement, EngagementFinding
from fraction.reporting.context import build_report_context
from fraction.reporting.render_docx import render_report_docx
from fraction.reporting.render_html import render_report_html


def _template(db, slug):
    return db.query(ChecklistTemplate).filter_by(slug=slug).one()


def _setup(session_factory):
    with session_factory() as db:
        e = Engagement(name="Report Eng", company_name="Acme")
        db.add(e)
        db.flush()
        finding = EngagementFinding(engagement_id=e.id, title="Egress leak to Internet")
        db.add(finding)
        db.flush()
        # coverage: mark one item failed + link the finding
        cov = C.assign_template(db, e, _template(db, "web-app-api"))
        cov.items[0].status = "fail"
        cov.items[0].finding_id = finding.id
        # compliance: PCI, mark a deficient + an na
        comp = C.assign_template(db, e, _template(db, "pci-dss-segmentation"))
        comp.items[0].status = "pass"
        comp.items[1].status = "fail"
        comp.items[2].status = "na"
        # reminder: stays internal (include_in_report defaults False)
        C.assign_template(db, e, _template(db, "global-pre-engagement"))
        db.commit()
        return e.id, finding.id


def test_report_context_excludes_reminder(session_factory):
    eid, _ = _setup(session_factory)
    with session_factory() as db:
        ctx = build_report_context(db.get(Engagement, eid))
    kinds = sorted(c.kind for c in ctx.checklists)
    assert kinds == ["compliance", "coverage"]  # reminder excluded by default


def test_report_context_toggle_includes_reminder(session_factory):
    eid, _ = _setup(session_factory)
    with session_factory() as db:
        e = db.get(Engagement, eid)
        rem = next(c for c in e.checklists if c.kind.value == "reminder")
        rem.include_in_report = True
        db.commit()
        ctx = build_report_context(db.get(Engagement, eid))
    assert "reminder" in {c.kind for c in ctx.checklists}


def test_report_html_renders_coverage_and_attestation(session_factory):
    eid, fid = _setup(session_factory)
    with session_factory() as db:
        html = render_report_html(build_report_context(db.get(Engagement, eid)))
    assert "Methodology and Coverage" in html
    assert "Compliance Attestation" in html
    assert "PCI-DSS" in html  # framework grouping heading
    assert "1.3.1" in html or "1.1" in html  # a control_ref appears
    # failed coverage item cross-links its finding by anchor
    assert f'href="#finding-{fid}"' in html
    assert "Egress leak to Internet" in html
    # attestation buckets rendered
    assert "ck-satisfied" in html and "ck-deficient" in html


def test_report_docx_renders_checklists(session_factory):
    from docx import Document

    eid, _ = _setup(session_factory)
    with session_factory() as db:
        data = render_report_docx(build_report_context(db.get(Engagement, eid)))
    doc = Document(io.BytesIO(data))
    text = "\n".join(p.text for p in doc.paragraphs)
    table_text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    assert "Methodology and Coverage" in text
    assert "Compliance Attestation" in text
    assert "PCI-DSS" in text  # framework heading
    # the compliance attestation table carries control refs + results
    assert "Control" in table_text and "Deficient" in table_text


def test_report_html_no_checklists_section_when_none(session_factory):
    with session_factory() as db:
        e = Engagement(name="Bare", company_name="X")
        db.add(e)
        db.commit()
        html = render_report_html(build_report_context(db.get(Engagement, e.id)))
    assert "Methodology and Coverage" not in html
    assert "Compliance Attestation" not in html
