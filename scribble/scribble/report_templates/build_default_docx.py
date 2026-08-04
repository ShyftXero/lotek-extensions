"""Authors ``scribble/report_templates/default.docx`` — the docxtpl template WS8 renders against.

Run this script to (re)generate the committed ``default.docx`` binary whenever the template layout
needs to change::

    uv run python -m scribble.report_templates.build_default_docx

Structure (mirrors the HTML report's shape, PLAN.md §9):

1. **Title page** — engagement/company name, client, dates, a CONFIDENTIAL marker, page break.
2. **Executive summary** — overall risk (color driven by a Jinja if/elif chain over
   ``rollup.overall``) + a severity-count table.
3. **Findings** — ``{% for group in groups %}`` … ``{% for f in group.findings %}``: finding title,
   a severity/CVSS/target meta table whose severity cell's background is set via docxtpl's
   ``{% cellbg <jinja ternary over f.severity> %}`` tag (this — not FACTION's `FAC701` sentinel hack —
   is "per-severity coloring via Jinja conditionals", PLAN.md §9), the finding's rich body
   (``{{r f.body }}``, filled by ``content/render_docx.py`` at render time), and an evidence list
   (``InlineImage`` per embeddable artifact, caption-only fallback otherwise).

This module only *authors* the template; it has no docxtpl/Jinja context of its own — every ``{{ }}``
/``{% %}`` string here is inert literal text until ``reporting/render_docx.py`` loads the resulting
``.docx`` as a ``docxtpl.DocxTemplate`` and calls ``.render()``.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUTPUT_PATH = Path(__file__).resolve().parent / "default.docx"

# Matches scribble/static/scribble.css's --sev-* ramp / render_html.py's CSS (kept in sync by hand —
# both are small, stable palettes) so the HTML and DOCX deliverables read the same.
SEVERITY_COLORS = {
    "critical": "B91C1C",
    "high": "DC2626",
    "medium": "EA580C",
    "low": "CA8A04",
    "info": "0284C7",
}
SEVERITY_ORDER = ("critical", "high", "medium", "low", "info")


def _severity_color_expr(var: str) -> str:
    """Build a nested Jinja ternary string mapping ``var`` (a severity string) to its hex color.

    Generated (not hand-typed) so the mapping can't drift out of sync with ``SEVERITY_COLORS`` /
    typo itself into a broken template.
    """
    expr = f'"{SEVERITY_COLORS[SEVERITY_ORDER[-1]]}"'
    for sev in reversed(SEVERITY_ORDER[:-1]):
        expr = f'("{SEVERITY_COLORS[sev]}" if {var} == "{sev}" else {expr})'
    return expr


def _tag_paragraph(doc: Document, text: str) -> None:
    """A paragraph whose sole content is literal Jinja markup (``{% for %}`` etc.) — renders to
    nothing once the template is filled."""
    p = doc.add_paragraph()
    p.add_run(text)


def _conditional_colored_run(paragraph, var: str, value_markup: str, color_map: dict[str, str]) -> None:
    """Emit an ``{% if %}/{% elif %}/…/{% else %}/{% endif %}`` chain, one differently-colored run
    of ``value_markup`` per branch — the "Jinja conditionals" per-severity coloring PLAN.md §9 calls
    for, used here for the executive-summary risk banner (the finding-level severity indicator uses
    the complementary ``{% cellbg %}`` cell-shading technique — see ``_add_finding_meta_table``)."""
    keys = list(color_map.keys())
    for i, key in enumerate(keys):
        kw = "if" if i == 0 else "elif"
        paragraph.add_run(f'{{% {kw} {var} == "{key}" %}}')
        run = paragraph.add_run(value_markup)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(color_map[key])
    paragraph.add_run("{% else %}")
    run = paragraph.add_run(value_markup)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("6B7280")
    paragraph.add_run("{% endif %}")


def _add_title_page(doc: Document) -> None:
    title = doc.add_paragraph(style="Title")
    title.add_run("{{ engagement_name }}")

    subtitle = doc.add_paragraph()
    subtitle.add_run("{{ company_name }}").bold = True

    meta = doc.add_paragraph()
    meta.add_run("Prepared for {{ client_name }}")

    dates = doc.add_paragraph()
    dates.add_run("{{ start_date }} – {{ end_date }}")

    scope = doc.add_paragraph()
    scope.add_run("Scope: {{ scope_type }}")

    generated = doc.add_paragraph()
    generated.add_run("Report generated {{ generated_date }}")

    confidential = doc.add_paragraph()
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = confidential.add_run("CONFIDENTIAL")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("B91C1C")

    doc.add_page_break()


def _add_executive_summary(doc: Document) -> None:
    doc.add_paragraph("Executive Summary", style="Heading 1")

    risk_p = doc.add_paragraph()
    risk_p.add_run("Overall Risk: ").bold = True
    _conditional_colored_run(
        risk_p,
        "rollup.overall",
        "{{ rollup.overall_label }}",
        {sev: SEVERITY_COLORS[sev] for sev in SEVERITY_ORDER},
    )

    summary_p = doc.add_paragraph()
    summary_p.add_run(
        "{{ groups|length }} section(s), {{ rollup.total }} finding(s) included in this report."
    )

    narrative_p = doc.add_paragraph()
    narrative_p.add_run("{{ narrative }}")

    table = doc.add_table(rows=2, cols=len(SEVERITY_ORDER) + 1)
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    data_cells = table.rows[1].cells
    for i, sev in enumerate(SEVERITY_ORDER):
        header_cells[i].paragraphs[0].add_run(sev.title()).bold = True
        data_cells[i].paragraphs[0].add_run(f"{{{{ rollup.counts.{sev} }}}}")
    header_cells[-1].paragraphs[0].add_run("Total").bold = True
    data_cells[-1].paragraphs[0].add_run("{{ rollup.total }}")

    doc.add_page_break()


def _add_finding_meta_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    sev_cell, cvss_cell, target_cell = table.rows[0].cells

    sev_p = sev_cell.paragraphs[0]
    sev_p.add_run(f"{{% cellbg {_severity_color_expr('f.severity')} %}}")
    sev_run = sev_p.add_run("{{ f.severity_label }}")
    sev_run.bold = True
    sev_run.font.color.rgb = RGBColor.from_string("FFFFFF")

    # Guarded (not unconditional runs): an empty ``f.cvss_score``/``f.target`` would otherwise render as
    # a bare "CVSS: "/"Target: " label with nothing after it — the shipped empty-chip bug (CONTRACT.md
    # §6 Track F). docxtpl merges a paragraph's run text before evaluating Jinja, so the ``{% if %}``/
    # ``{% endif %}`` tags can live in their own runs either side of the label run.
    cvss_p = cvss_cell.paragraphs[0]
    cvss_p.add_run("{% if f.cvss_score %}")
    cvss_p.add_run("CVSS: {{ f.cvss_score }}")
    cvss_p.add_run("{% endif %}")

    target_p = target_cell.paragraphs[0]
    target_p.add_run("{% if f.target %}")
    target_p.add_run("Target: {{ f.target }}")
    target_p.add_run("{% endif %}")


def _add_evidence_section(doc: Document) -> None:
    _tag_paragraph(doc, "{% if f.artifacts %}")
    doc.add_paragraph("Evidence", style="Heading 4")
    _tag_paragraph(doc, "{% for a in f.artifacts %}")
    _tag_paragraph(doc, "{% if a.embedded %}")
    image_p = doc.add_paragraph()
    image_p.add_run("{{ a.image }}")
    cap = doc.add_paragraph()
    cap.add_run("{{ a.caption }}").italic = True
    _tag_paragraph(doc, "{% else %}")
    missing_p = doc.add_paragraph()
    missing_p.add_run("\U0001f4c4 {{ a.filename }} — {{ a.caption }} (not embedded)").italic = True
    _tag_paragraph(doc, "{% endif %}")
    _tag_paragraph(doc, "{% endfor %}")
    _tag_paragraph(doc, "{% endif %}")


def _add_findings_body(doc: Document) -> None:
    _tag_paragraph(doc, "{% if not groups %}")
    doc.add_paragraph("No findings recorded for this engagement.")
    _tag_paragraph(doc, "{% endif %}")

    _tag_paragraph(doc, "{% for group in groups %}")
    doc.add_paragraph("{{ group.name }}", style="Heading 1")

    _tag_paragraph(doc, "{% if not group.findings %}")
    doc.add_paragraph("No findings in this section.")
    _tag_paragraph(doc, "{% endif %}")

    _tag_paragraph(doc, "{% for f in group.findings %}")
    doc.add_paragraph("{{ f.title }}", style="Heading 2")
    _add_finding_meta_table(doc)
    body_p = doc.add_paragraph()
    body_p.add_run("{{r f.body }}")
    _add_evidence_section(doc)
    _tag_paragraph(doc, "{% endfor %}")  # f in group.findings

    _tag_paragraph(doc, "{% endfor %}")  # group in groups


def build() -> Document:
    doc = Document()
    _add_title_page(doc)
    _add_executive_summary(doc)
    _add_findings_body(doc)
    return doc


def main() -> None:
    doc = build()
    doc.save(str(OUTPUT_PATH))
    print(f"wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
