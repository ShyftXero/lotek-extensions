"""Editable ``.docx`` report renderer (WS8).

Consumes ``ReportContext`` (the FROZEN contract in ``scribble.reporting.context``) only — no DB
access, no Flask — mirroring WS7's HTML renderer. Loads ``report_templates/default.docx`` (a docxtpl
template authored by ``report_templates/build_default_docx.py``) and fills it with:

- Title-page + executive-summary scalars (company/engagement name, dates, severity rollup).
- Nested ``{% for group in groups %}`` / ``{% for f in group.findings %}`` loops (board order ==
  document order, exactly as the context already orders things).
- Each finding's content blocks (``FindingCtx.blocks_html`` — already variable-resolved, sanitized
  HTML) converted to a single :class:`docxtpl.RichText` via ``content/render_docx.py``'s HTML walker,
  bound to the template's ``{{r f.body }}`` field.
- Evidence-gallery artifacts embedded via :class:`docxtpl.InlineImage` (skipped gracefully — caption
  text only — when bytes aren't available or the artifact isn't an image).
- Per-severity coloring authored as Jinja conditionals *in the template itself*
  (``report_templates/build_default_docx.py``'s ``{% cellbg ... %}`` cell-shading expression) —
  replacing FACTION's `FAC701` sentinel hack per PLAN.md §9.

Inline (``inlineImage``/``figure``) images embedded *inside* prose round-trip through the same
placeholder-URL trick WS7 uses (see ``render_html.make_inline_artifact_url``), but with WS8's own
prefix so this module never depends on WS7's private internals: callers build the ``ReportContext``
with ``artifact_url=`` wired to :func:`make_inline_artifact_url` from *this* module for those images
to be embeddable here.
"""

from __future__ import annotations

import io
import json
import re
from collections.abc import Callable
from dataclasses import dataclass
from datetime import UTC, datetime
from html import escape as _html_escape
from pathlib import Path
from urllib.parse import quote, unquote

from docx.shared import Mm
from docxtpl import DocxTemplate, InlineImage, RichText

from scribble.content.render_docx import html_to_richtext
from scribble.enums import SEVERITY_ORDER as _ENUM_SEVERITY_ORDER
from scribble.reporting.context import (
    DIAGRAM_CAPTION_FALLBACK,
    ArtifactCtx,
    FindingCtx,
    GroupCtx,
    ReportContext,
    figure_caption,
)

ArtifactBytes = Callable[[str], "bytes | None"]

SEVERITY_ORDER: tuple[str, ...] = tuple(s.value for s in _ENUM_SEVERITY_ORDER)  # worst-first

_TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "report_templates" / "default.docx"

_BLOCK_LABELS = {"description": "Description", "remediation": "Remediation", "details": "Details"}
_BLOCK_ORDER = ("description", "remediation", "details")

# Same placeholder trick as ``render_html.make_inline_artifact_url`` (schemeless + relative so it
# survives ``content/render_html.py``'s sanitize pass), but WS8-owned so this module never has to
# reach into WS7's private constants.
_INLINE_PREFIX = "/__scribble_docx_inline__/"

_EVIDENCE_IMAGE_WIDTH = Mm(70)

# Don't embed an evidence image whose bytes exceed this — fall back to the caption-only rendering — so
# a single huge artifact can't blow up the render's memory footprint (matches the cap in
# ``content/render_docx.py`` and the stat-based guard in ``report_docx_api``). 25 MiB.
_MAX_EVIDENCE_BYTES = 25 * 1024 * 1024


def make_inline_artifact_url(storage_path: str | None) -> str:
    """Placeholder ``src`` for an inline content-node image, to pass as (or wire into) the
    ``artifact_url`` callback given to ``build_report_context`` for WS8 rendering."""
    if not storage_path:
        return ""
    return _INLINE_PREFIX + quote(storage_path, safe="")


def _decode_inline_placeholder(src: str | None) -> str | None:
    if not src or not src.startswith(_INLINE_PREFIX):
        return None
    return unquote(src[len(_INLINE_PREFIX) :]) or None


def _make_image_resolver(artifact_bytes: ArtifactBytes | None) -> Callable[[str], bytes | None]:
    def _resolve(src: str) -> bytes | None:
        if not artifact_bytes:
            return None
        storage_path = _decode_inline_placeholder(src) or src
        try:
            return artifact_bytes(storage_path)
        except Exception:
            return None

    return _resolve


def _label_run_xml(text: str) -> str:
    rt = RichText()
    rt.add(text)
    return rt.xml


def _child_host_label(c: FindingCtx) -> str:
    if c.target_host:
        return c.target_host + (f":{c.target_port}" if c.target_port else "")
    if c.target_url:
        return c.target_url
    return c.title


def _child_summary_text(c: FindingCtx) -> str:
    """The per-host evidence line for a child finding's "Affected Hosts" list entry — mirrors
    ``render_html._child_summary_text``: built from the child's OWN ``variables`` overlay
    (``FindingCtx.facts_line``), not from its content blocks, since every child promoted under the same
    vuln-DB template shares that template's ``content_json`` verbatim with its parent."""
    return c.facts_line


def _children_html(children: list[FindingCtx]) -> str:
    """A small HTML fragment (fed through the same ``html_to_richtext`` walker as content blocks) —
    the "Affected Hosts" compact per-host list, rendered INLINE into the parent finding's body since
    the ``.docx`` template's finding loop is authored as a flat paragraph sequence
    (``build_default_docx.py``), not a per-finding sub-scope a nested ``{%tr for %}`` row-loop could
    cleanly hang off of. This is the DOCX-side mirror of ``render_html._render_children``'s
    ``<details>``/table.

    A child's OWN artifacts render here too (issue #54, the docx half of ext#40 mechanism 2): before
    this, a screenshot attached to a promoted per-host instance never appeared in the docx report at
    all -- exactly the render_html gap ``_render_child_evidence_cell`` fixed on the HTML side. Each
    artifact becomes a placeholder ``<img>`` (this module's OWN ``make_inline_artifact_url``/
    ``_INLINE_PREFIX``, resolved by the ``image_resolver`` already threaded into ``html_to_richtext``
    -- see ``_finding_body_richtext``); a non-image or an image whose bytes are unavailable/oversized
    degrades to ``content/render_docx.py``'s existing bracketed placeholder rather than the render
    failing."""
    items = []
    for c in children:
        host = _html_escape(_child_host_label(c))
        summary = _child_summary_text(c)
        if summary:
            items.append(f"<li><strong>{host}</strong> — {_html_escape(summary)}</li>")
        else:
            items.append(f"<li><strong>{host}</strong></li>")
        for a in c.artifacts:
            src = make_inline_artifact_url(a.storage_path)
            if not src:
                continue
            cap = _html_escape(_numbered_caption(a))  # ext#117 — same "Figure N" the HTML prints
            items.append(f'<li><img src="{src}" alt="{cap}"/> {cap}</li>')
    return f"<h4>Affected Hosts ({len(children)})</h4><ul>{''.join(items)}</ul>"


def _metadata_line_html(f: FindingCtx) -> str:
    """A DOCX-side "Classification" line mirroring render_html's header chips (#625). Word's binary
    finding template has no place to add header cells, so — exactly as #620's risk-override note did —
    the metadata is folded into the finding BODY as an omit-when-empty leading line rather than
    regenerating ``default.docx``. KEV/EPSS carry the snapshot ``as_of`` so they never assert a stale
    fact as current."""
    parts: list[str] = []
    if f.cwe_ids:
        parts.append("CWE: " + ", ".join(f.cwe_ids))
    if f.cve_ids:
        parts.append("CVE: " + ", ".join(f.cve_ids))
    if f.owasp_categories:
        parts.append("OWASP: " + ", ".join(f.owasp_categories))
    ti = f.threat_intel
    if ti:
        as_of = ti.get("as_of")
        suffix = f" (as of {as_of})" if as_of else ""
        if ti.get("kev"):
            parts.append(f"KEV{suffix}")
        epss = ti.get("epss")
        if isinstance(epss, (int, float)):
            parts.append(f"EPSS {epss:.2f}{suffix}")
    if not parts:
        return ""
    return f"<p><strong>Classification:</strong> {_html_escape(' · '.join(parts))}</p>"


def _references_html(f: FindingCtx) -> str:
    """The DOCX References list (#624): non-suppressed refs as labeled hyperlinks (``html_to_richtext``
    renders ``<a href>`` as a real Word hyperlink for a safe http(s)/mailto scheme), OMITTED when there
    are none. Mirror of ``render_html._render_references``; ``source`` is not shown (#624 Q4)."""
    items: list[str] = []
    for r in f.references:  # already filtered to non-suppressed in context._finding_ctx
        label = str(r.get("label") or r.get("url") or "").strip()
        if not label:
            continue
        url = str(r.get("url") or "").strip()
        if url:
            items.append(f'<li><a href="{_html_escape(url)}">{_html_escape(label)}</a></li>')
        else:
            items.append(f"<li>{_html_escape(label)}</li>")
    if not items:
        return ""
    return f"<h4>References</h4><ul>{''.join(items)}</ul>"


def _finding_body_richtext(
    tpl: DocxTemplate,
    blocks_html: dict[str, str],
    image_resolver: Callable[[str], bytes | None],
    *,
    children: list[FindingCtx] | None = None,
    metadata_html: str = "",
    references_html: str = "",
) -> RichText:
    """Combine a finding's content blocks into one RichText, each under a "Heading 3" label —
    mirrors ``render_html._render_blocks``'s label + ordering behavior. ``children`` (a nested
    finding's per-host instances) render as an "Affected Hosts" list appended to the same body.
    ``metadata_html`` (a leading Classification line, #625) and ``references_html`` (a trailing
    References list, #624) are omit-when-empty HTML fragments fed through the same richtext walker."""
    parts: list[str] = []
    any_open = False

    def _open(style: str | None = None) -> None:
        nonlocal any_open
        prefix = "</w:p><w:p>" if any_open else ""
        ppr = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
        parts.append(prefix + ppr)
        any_open = True

    if metadata_html:
        _open()
        meta_rt = html_to_richtext(metadata_html, tpl=tpl, image_resolver=image_resolver)
        if meta_rt.xml:
            parts.append(meta_rt.xml)

    # ``references`` render via the structured block below (references_html). Suppress a legacy prose
    # ``references`` content block ONLY when there ARE structured refs (else double-render); when the
    # column is empty, let the legacy block render as before -- an existing finding that stored refs as a
    # prose block (pre-#624) must not silently lose them from the DOCX (no migration backfills). (#624)
    seen: set[str] = {"references"} if references_html else set()
    ordered_keys = [k for k in _BLOCK_ORDER] + [k for k in blocks_html if k not in _BLOCK_ORDER]
    rendered_any = False
    for key in ordered_keys:
        if key in seen:
            continue
        seen.add(key)
        fragment = blocks_html.get(key)
        if not fragment:
            continue
        rendered_any = True
        label = _BLOCK_LABELS.get(key, key.replace("_", " ").title())
        _open(style="Heading3")  # styleId (spaceless) — <w:pStyle> resolves by id, not display name
        parts.append(_label_run_xml(label))
        block_rt = html_to_richtext(fragment, tpl=tpl, image_resolver=image_resolver)
        if block_rt.xml:
            parts.append("</w:p><w:p>")
            parts.append(block_rt.xml)
            any_open = True

    if not rendered_any and not children:
        _open()
        parts.append(_label_run_xml("No content."))

    if children:
        if not any_open:
            _open()
        children_rt = html_to_richtext(_children_html(children), tpl=tpl, image_resolver=image_resolver)
        if children_rt.xml:
            parts.append("</w:p><w:p>")
            parts.append(children_rt.xml)
            any_open = True

    if references_html:
        if not any_open:
            _open()
        refs_rt = html_to_richtext(references_html, tpl=tpl, image_resolver=image_resolver)
        if refs_rt.xml:
            parts.append("</w:p><w:p>")
            parts.append(refs_rt.xml)
            any_open = True

    rt = RichText()
    rt.xml = "".join(parts)
    return rt


def _target_text(f: FindingCtx) -> str:
    bits: list[str] = []
    if f.target_host:
        bits.append(f.target_host + (f":{f.target_port}" if f.target_port else ""))
    if f.target_url:
        bits.append(f.target_url)
    return " · ".join(bits)


def _numbered_caption(a: ArtifactCtx) -> str:
    """``"Figure 3 — Payload firing in the browser"`` (ext#117). The number comes off the CONTEXT
    (``context.number_figures``), never from a counter this renderer keeps, so it is the same number
    the HTML deliverable prints for the same artifact."""
    return _xml_safe(figure_caption(a.figure_number, a.caption or a.filename))


def _artifact_ctx(
    a: ArtifactCtx, artifact_bytes: ArtifactBytes | None, tpl: DocxTemplate
) -> dict[str, object]:
    is_image = (a.content_type or "").startswith("image/")
    image = None
    if is_image and artifact_bytes:
        try:
            data = artifact_bytes(a.storage_path)
        except Exception:
            data = None
        if data and len(data) > _MAX_EVIDENCE_BYTES:
            data = None  # oversized: fall back to caption-only rather than embed a huge blob
        if data:
            try:
                image = InlineImage(tpl, io.BytesIO(data), width=_EVIDENCE_IMAGE_WIDTH)
            except Exception:
                image = None
    return {
        "caption": _numbered_caption(a),
        "filename": a.filename,
        "image": image,
        "embedded": image is not None,
    }


def _finding_ctx(
    f: FindingCtx, *, tpl: DocxTemplate, artifact_bytes: ArtifactBytes | None
) -> dict[str, object]:
    image_resolver = _make_image_resolver(artifact_bytes)
    sev = f.severity if f.severity in SEVERITY_ORDER else "info"
    return {
        "title": f.title,
        "severity": sev,
        "severity_label": sev.title(),
        # lotek#618: the same client-facing label the HTML badge shows, from the same predicate.
        # Empty for a `new` finding, and the template's paragraph-level `{% if %}` then drops the
        # whole line — the DOCX is a peer deliverable, not a lossy copy of the HTML.
        "status_label": f.status_label,
        "disposition": f.disposition,
        "cvss_score": f"{f.cvss_score:.1f}" if f.cvss_score is not None else "",
        "cvss_vector": f.cvss_vector or "",
        "target": _target_text(f),
        "body": _finding_body_richtext(
            tpl, f.blocks_html, image_resolver, children=f.children,
            metadata_html=_metadata_line_html(f), references_html=_references_html(f),
        ),
        "artifacts": [_artifact_ctx(a, artifact_bytes, tpl) for a in f.artifacts],
    }


def _group_ctx(g: GroupCtx, *, tpl: DocxTemplate, artifact_bytes: ArtifactBytes | None) -> dict:
    return {
        "name": g.name,
        "type_slug": g.type_slug or "",
        "findings": [_finding_ctx(f, tpl=tpl, artifact_bytes=artifact_bytes) for f in g.findings],
    }


def _build_context(ctx: ReportContext, *, tpl: DocxTemplate, artifact_bytes: ArtifactBytes | None) -> dict:
    rollup = ctx.rollup
    counts = rollup.counts if rollup else {}
    computed = rollup.overall if rollup else "info"
    total = rollup.total if rollup else 0
    clean = total == 0
    computed_label = "No Findings" if clean else computed.title()
    # lotek#620: an operator override is an authored judgement layered on the COMPUTED band. Word has no
    # banner CSS to restyle and the banner is authored into the binary template, so the marker + original
    # computed band are baked into ``overall_label`` and the rationale is appended to the summary
    # narrative as an attributed, present-tense sentence — parity with the HTML banner (marker, computed
    # value, rationale) without regenerating ``default.docx``. ``overall`` (used for cell colour) follows
    # the effective band so the verdict card colour matches the headline.
    override = ctx.risk_override
    effective = override or computed
    if override:
        overall = effective
        overall_label = f"{effective.title()} — assessor-adjusted (computed: {computed_label})"
    else:
        overall = computed
        overall_label = computed_label
    narrative = ctx.narrative or ""
    if override:
        note = (
            f"The overall risk rating shown is the assessor's adjustment of the computed "
            f"{computed_label} rating to {effective.title()}."
        )
        rationale = (ctx.risk_override_rationale or "").strip()
        if rationale:
            note += f" Rationale: {rationale}"
        narrative = f"{narrative} {note}".strip() if narrative else note
    return {
        "company_name": ctx.company_name or "",
        "engagement_name": ctx.engagement_name,
        "client_name": ctx.client_name or "",
        "scope_type": ctx.scope_type or "",
        "start_date": ctx.start_date or "",
        "end_date": ctx.end_date or "",
        "generated_date": datetime.now(UTC).strftime("%Y-%m-%d"),
        "narrative": narrative,
        "rollup": {
            "counts": {s: counts.get(s, 0) for s in SEVERITY_ORDER},
            "total": total,
            "overall": overall,
            "overall_label": overall_label,
        },
        "groups": [_group_ctx(g, tpl=tpl, artifact_bytes=artifact_bytes) for g in ctx.groups],
    }


# ── attack paths (ext#115) ───────────────────────────────────────────────────────────────────────────
#
# The HTML deliverable embeds vector's self-contained ``export.html`` in a sandboxed iframe and the
# animation plays. Word has no browser, so until ext#115 the .docx simply had no attack path in it at
# all: `Attack path` did not appear in `word/document.xml`, and a reader given only the .docx could not
# tell that a diagram existed. That is a SILENT content drop between two deliverables of the same
# engagement, which is worse than an ugly rendition.
#
# Why a native Word table and not a picture: a picture in a .docx must be RASTER (python-docx/docxtpl
# reject SVG; Word's `svgBlip` needs a PNG fallback anyway), and vector's diagram only exists as pixels
# once a browser has run `vector-viewer.js`. Rasterizing it server-side means shipping a headless
# browser or a rasterizer into a mounted production extension; re-drawing it in Python means a SECOND
# renderer that drifts from the JS the first time anyone edits the viewer. So this draws the same
# geometry the viewer draws -- `zone` is the column, `row` is the row (vector-viewer.js `geometry()`) --
# with Word's own table, plus the phase walkthrough and the edge list. Static, selectable, searchable,
# and it cannot drift out of a font metric.
#
# ponytail: table rendition, not pixels. Upgrade path if a true still is ever wanted: have vector's
# viewer serialize its live <svg> at export time (`XMLSerializer`, ~5 lines, zero drift) and store it
# alongside `embed_html`, then rasterize THAT here.

# ``embed_html`` is operator/agent-supplied (it arrives over a PAT POST -- see
# ``api_pat.scribble_link_attack_path``) and is stored verbatim, so nothing guarantees it came from
# vector or is bounded the way vector's own schema bounds a model. Every number below is scribble's own
# cap on how much of a snapshot this renderer will turn into document, independent of vector's limits.
# Set to ``api_pat._MAX_DIAGRAM_HTML_BYTES``'s number (10 MiB) so nothing the link route ACCEPTS is
# stored and then permanently unrenderable in Word. Note the units differ -- that cap counts UTF-8
# BYTES, this one counts CHARACTERS, and ``len(str) <= len(utf8)`` -- so this bound is strictly the
# looser of the two and never binds a row the route wrote. It bites a row written another way. Safe to
# raise now the scan is linear: 8 MiB measured 0.02s (the regex this replaced would have taken hours).
_MAX_DIAGRAM_SCAN_CHARS = 10 * 1024 * 1024
# ...and the same bound again for the whole SECTION. ``_MAX_DIAGRAM_SCAN_CHARS`` bounds one snapshot;
# nothing bounded the product, and the two caps multiply: 50 linked diagrams of 9.6 MiB each measured
# 34s of GIL-held CPU and 515 MiB peak per ``GET …/report?format=docx``, from rows any ``write``-scope
# PAT can store and any ``read``-scope viewer can then trigger, repeatedly. A budget spent across the
# section makes the section's cost independent of how many diagrams are linked; a diagram that arrives
# after it is exhausted degrades to the same honest "interactive figure in the HTML report" note the
# renderer already emits for a snapshot it cannot read.
_MAX_REPORT_SCAN_CHARS = 10 * 1024 * 1024
_MAX_DIAGRAMS = 50  # a report section, not an archive (cf. render_html's _MAX_APPENDIX_ITEMS)
_MAX_DIAGRAM_ZONES = 12  # a Word table wider than this is unreadable on a portrait page anyway
_MAX_DIAGRAM_ROWS = 40
_MAX_DIAGRAM_NODES = 2000
_MAX_DIAGRAM_PHASES = 60
_MAX_DIAGRAM_EDGES = 80
_MAX_DIAGRAM_TEXT = 400  # per field, matching vector's own _MED/_LONG spirit

# Codepoints legal in a Python str and legal in JSON, but ILLEGAL in XML 1.0 -- lxml (under
# python-docx) raises on any of them, which would turn one poisoned diagram into an uncaught 500 on
# every ``.docx`` export of that engagement, forever. They cannot be filtered at the link route: the
# JSON blob carries them as the six ASCII characters ``\u0000``, so the stored ``embed_html`` holds no
# literal control byte for ``api_pat._nul_safe`` to strip, and ``json.loads`` materializes the real
# character here. (Tab/LF/CR are legal in XML and are collapsed by ``_d_str``'s whitespace split
# anyway.)
#
# THREE classes, not one -- the first pass only covered C0 and a later review found the other two
# still live, each with a working payload:
#   * C0 controls                -> ValueError("All strings must be XML compatible")
#   * lone surrogates D800-DFFF  -> UnicodeEncodeError out of lxml's utf-8 encode
#   * the noncharacters FFFE/FFFF-> ValueError, same as C0
# C1 (80-9F), FDD0-FDEF and the non-BMP planes were measured and are fine; do not widen further, a
# scrub that eats legible text is its own defect.
_XML_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f\ud800-\udfff\ufffe\uffff]")


def _xml_safe(text: str) -> str:
    """Drop the characters python-docx cannot serialize. Applied at the TWO places untrusted text
    becomes document -- here and in :func:`_numbered_caption` -- not at each ``add_run`` call."""
    return _XML_ILLEGAL_RE.sub("", text)


def _d_str(value: object, cap: int = _MAX_DIAGRAM_TEXT) -> str:
    """A capped, single-line, XML-safe string from an untrusted model field; non-scalar -> ``""``."""
    if value is None or isinstance(value, bool) or not isinstance(value, (str, int, float)):
        return ""
    # Slice BEFORE normalising: whitespace-splitting a 5 MiB field cost 0.11s and 26 MiB of peak RSS to
    # produce 400 characters. ``cap * 8`` leaves room for a run of separators to collapse away without
    # the result coming up short.
    return " ".join(_xml_safe(str(value)[: cap * 8]).split())[:cap]


def _d_int(value: object, default: int = 0) -> int:
    try:
        if isinstance(value, bool):
            return default
        return int(value)
    except (TypeError, ValueError, OverflowError):
        # OverflowError is an ArithmeticError, NOT a ValueError: ``int(float("inf"))`` raises it, and
        # JSON's non-standard ``Infinity`` / an overflowing ``1e999`` both produce that float. Missing
        # it made a one-token payload an uncaught 500 on the whole deliverable.
        return default


def _d_list(value: object) -> list:
    return value if isinstance(value, list) else []


def _find_model_blob(embed_html: str) -> str | None:
    """The text inside ``<script … id="vap-model">…</script>``, by LINEAR scans.

    This was a regex (``<script\\b[^>]*\\bid=["']vap-model["'][^>]*>(.*?)</script>``) and that was a
    denial of service: two unanchored ``[^>]*`` runs before a literal mean every ``<script`` in the
    input is a start position costing O(n), so the match is O(n²) in a string an operator PAT chooses.
    Measured on the real pattern: ``"<script " * 8000`` (64 KiB) took **15.9 s**, and the curve
    extrapolates to ~71 minutes at 1 MiB -- well inside the 10 MiB the link route accepts. CPython's
    ``re`` holds the GIL and never yields to the gevent hub, so that is not a slow request, it is the
    whole worker. Lowering the size cap does not fix a quadratic; removing the backtracking does.
    (Re-measured after: 8 MiB of the same payload, 0.02 s.)

    Deliberately strict about what it accepts, because the job is to read VECTOR's output, not to
    parse hostile HTML: the id attribute must sit inside a ``<script`` tag that opens within 256
    characters before it, and the body ends at the first ``</script``. vector's ``json_for_script``
    escapes ``<``/``>``/``&`` as ``\\uXXXX``, so a genuine model never contains either."""
    for needle in ('id="vap-model"', "id='vap-model'"):
        at = embed_html.find(needle)
        while at >= 0:
            start = embed_html.rfind("<script", max(0, at - 256), at)
            body = embed_html.find(">", at) if start >= 0 else -1
            end = embed_html.find("</script", body + 1) if body >= 0 else -1
            if end >= 0:
                return embed_html[body + 1 : end]
            # Not a real ``<script id="vap-model">`` -- keep scanning. Bailing on the FIRST hit let a
            # decoy (``<!-- id="vap-model" -->``) hide the genuine model, so Word printed "static
            # rendition not available" while the HTML iframe drew the diagram in full. Still linear:
            # each pass advances past the needle it just rejected.
            at = embed_html.find(needle, at + len(needle))
    return None


def _diagram_model(embed_html: str | None, budget: list[int] | None = None) -> dict | None:
    """The ``vector.attackpath/v1`` document carried inside a stored snapshot, or ``None``.

    Reads a JSON blob out of HTML scribble already stores -- it does not import vector (extensions stay
    independent -- CLAUDE.md) and does not execute anything. Never raises: a snapshot that is not
    vector's, is truncated, or is not JSON degrades to ``None`` and the caller still emits the section
    heading + caption, so the diagram is never silently absent."""
    if not embed_html or len(embed_html) > _MAX_DIAGRAM_SCAN_CHARS:
        return None
    if budget is not None:
        # One mutable cell shared across the section (see ``_MAX_REPORT_SCAN_CHARS``). Charged BEFORE
        # the scan, so an over-budget snapshot is never walked at all.
        if budget[0] < len(embed_html):
            return None
        budget[0] -= len(embed_html)
    blob = _find_model_blob(embed_html)
    if blob is None:
        return None
    try:
        # ``parse_constant`` kills JSON's non-standard ``Infinity``/``-Infinity``/``NaN`` at the parser
        # rather than leaving every downstream ``int()`` to survive a float special.
        model = json.loads(blob, parse_constant=lambda _name: None)
    except (ValueError, RecursionError):
        return None
    return model if isinstance(model, dict) else None


def _styled_paragraph(doc, style: str):
    """``doc.add_paragraph(style=...)``, degrading to an unstyled paragraph if the template lacks the
    style. A missing style is a KeyError from python-docx, and losing the whole deliverable over a
    cosmetic style would be absurd."""
    try:
        return doc.add_paragraph(style=style)
    except KeyError:
        return doc.add_paragraph()


# The status chip, as ``static/vector-viewer.js``'s ``nodeVisual()`` computes it at the final keyframe.
# Its rule is three-tiered and the order matters -- an earlier pass here got all three tiers wrong:
#   1. an explicit ``states[].label`` wins OUTRIGHT (last one wins), printed VERBATIM -- not uppercased;
#   2. otherwise the highest-``precedence`` state whose key is IN the catalog, printed as the catalog's
#      display label (``impacted`` shows "IMPACT", not "IMPACTED"). A key that is NOT in the catalog is
#      not a candidate at all -- the viewer shows nothing for it, so neither do we;
#   3. otherwise the node's ROLE supplies the chip (``roles[role].status``), or its idle text when the
#      role is idle and the node never activates.
# ``node.context`` nodes carry no chip at all.
#
# ponytail: the two BUILT-IN catalogs only. A model may override ``style.nodeStates``/``style.roles``
# (``vector-viewer.js``'s ``mergeStyle``), and a styled diagram's chips will read from the built-in
# label instead. Mirroring the merge would mean re-implementing the style resolver; upgrade path is to
# have vector serialize the resolved chip text into the model at export time.
_NODE_STATES = {
    "target": (1, "TARGET"),
    "owned": (3, "OWNED"),
    "beacon": (3, "BEACON"),
    "impacted": (4, "IMPACT"),
}
_NODE_ROLES = {
    "c2": ("C2", None),
    "rshell": ("REV SHELL", None),
    "stager": ("STAGER", None),
    "payload": ("PAYLOAD", None),
    "egress": ("EGRESS", "\u2014"),
    "backup": ("BACKUP", "STANDBY"),
}


def _note_truncation(doc, model: dict, drawn: _Drawn | None) -> None:
    """Name anything this renderer dropped from the model. Scribble's caps are TIGHTER than vector's
    (12 zones vs 40, 60 phases vs 200, 80 edges vs 1500, 2000 nodes vs 600-per-vector), so a large
    genuine diagram can lose content here -- and a silently truncated figure is the same defect class
    ext#115 is about.

    Counted against what the grid ACTUALLY DREW, not against the caps. Counting against the caps
    under-reported, and it under-reported the hosts specifically: a node is also dropped when its zone
    fell past the zone cap, when its zone id is unknown, when two zones shared an id and de-duped, or
    when it is not a dict at all. On the branch's own 400-zone fixture the old arithmetic announced
    "388 zones" and said nothing about the 388 hosts that vanished with them."""
    dropped = []
    zones_declared = len(_d_list(model.get("zones")))
    nodes_declared = len(_d_list(model.get("nodes")))
    for label, over in (
        ("zone", zones_declared - (drawn.zones if drawn else 0)),
        ("host", nodes_declared - (len(drawn.node_ids) if drawn else 0)),
        ("phase", len(_d_list(model.get("phases"))) - _MAX_DIAGRAM_PHASES),
        ("connection", len(_d_list(model.get("edges"))) - _MAX_DIAGRAM_EDGES),
    ):
        if over > 0:
            dropped.append(f"{over} {label}{'s' if over != 1 else ''}")
    if drawn and drawn.rows_clamped:
        # Not a drop -- a COLLAPSE. Those hosts are in the table, stacked into its last row rather than
        # placed where the viewer places them, which is worth saying out loud for the same reason.
        dropped.append(f"the placement of {drawn.rows_clamped} host"
                       f"{'s' if drawn.rows_clamped != 1 else ''} below row {_MAX_DIAGRAM_ROWS}")
    if not dropped:
        return
    note = doc.add_paragraph()
    note.add_run(
        "This rendition omits " + ", ".join(dropped)
        + " beyond what a page-width table can carry; the HTML report shows the diagram in full."
    ).italic = True


def _node_state_label(node: dict) -> str:
    """The node's status chip at the FINAL keyframe, chosen the way ``nodeVisual()`` chooses it -- see
    the tiers documented on :data:`_NODE_STATES`."""
    if node.get("context"):
        return ""
    explicit, best_rank, best_key = "", -1, ""
    for st in _d_list(node.get("states")):
        if not isinstance(st, dict):
            continue
        key = _d_str(st.get("state"), 40).lower()
        if key in _NODE_STATES:
            rank = _NODE_STATES[key][0]
            if rank >= best_rank:
                best_rank, best_key = rank, key
        label = _d_str(st.get("label"), 40)
        if label:
            explicit = label  # last one wins, exactly as the viewer's forEach does
    if explicit:
        return explicit
    if best_key:
        return _NODE_STATES[best_key][1]
    status, idle_status = _NODE_ROLES.get(_d_str(node.get("role"), 40).lower(), ("", None))
    if idle_status is not None and node.get("activateAt") is None:
        # An idle role stays idle unless the node activates; at the final keyframe an ``activateAt``
        # has always been reached, so its presence alone is enough.
        return idle_status
    return status


def _node_cell_text(node: dict) -> str:
    lines = [_d_str(node.get("label")) or _d_str(node.get("id"))]
    addr = _d_str(node.get("ip")) or _d_str(node.get("domain"))
    if addr:
        lines.append(addr)
    state = _node_state_label(node)
    if state:
        lines.append(f"[{state}]")
    return "\n".join(x for x in lines if x)


@dataclass(frozen=True)
class _Drawn:
    """What :func:`_add_diagram_grid` actually put in the table, so the truncation note and the edge
    list can both be honest about it rather than re-deriving it from the caps."""

    zones: int
    node_ids: set[str]
    rows_clamped: int


def _add_diagram_grid(doc, model: dict) -> _Drawn | None:
    """The static frame: zones as columns, ``row`` as rows — the same placement ``vector-viewer.js``'s
    ``geometry()`` computes. Returns ``None`` when the model has nothing placeable to draw."""
    zones = [z for z in _d_list(model.get("zones")) if isinstance(z, dict)]
    zones.sort(key=lambda z: _d_int(z.get("order")))
    # De-dupe by id BEFORE capping: two zones sharing an id would otherwise key the same bucket and
    # render identical columns. vector's ``geometry()`` is last-wins over the same map; first-wins here
    # matches the model's own declared order, and either way one id is one column.
    seen_ids: set[str] = set()
    zones = [z for z in zones if not (_d_str(z.get("id")) in seen_ids or seen_ids.add(_d_str(z.get("id"))))]
    zones = zones[:_MAX_DIAGRAM_ZONES]
    if not zones:
        return None
    by_zone: dict[str, dict[int, list[dict]]] = {_d_str(z.get("id")): {} for z in zones}
    max_row = 0
    placed: set[str] = set()
    rows_clamped = 0
    # ``nodes`` is the one list vector's own caps do not reach here, and every node sharing a
    # ``(zone, row)`` is concatenated into ONE cell -- one ``<w:br/>`` element each. Measured
    # uncapped: 100k nodes = 8.8s and 204 MB peak per render, multiplied by however many diagrams are
    # linked. Cap it like the rest.
    for node in _d_list(model.get("nodes"))[:_MAX_DIAGRAM_NODES]:
        if not isinstance(node, dict):
            continue
        zone_id = _d_str(node.get("zone"))
        if zone_id not in by_zone:
            continue  # a node in a zone this table does not show (unknown id, or past the zone cap)
        declared_row = _d_int(node.get("row"))
        row = max(0, min(declared_row, _MAX_DIAGRAM_ROWS - 1))
        if declared_row > row:
            rows_clamped += 1
        by_zone[zone_id].setdefault(row, []).append(node)
        placed.add(_d_str(node.get("id")))
        max_row = max(max_row, row)
    if not any(by_zone.values()):
        return None

    table = doc.add_table(rows=1, cols=len(zones))
    table.style = "Table Grid"
    for cell, zone in zip(table.rows[0].cells, zones, strict=True):
        cell.paragraphs[0].add_run(_d_str(zone.get("title")) or _d_str(zone.get("id"))).bold = True
        subtitle = _d_str(zone.get("subtitle"))
        if subtitle:
            cell.add_paragraph(subtitle)
    for row_idx in range(max_row + 1):
        cells = table.add_row().cells
        for cell, zone in zip(cells, zones, strict=True):
            nodes = by_zone[_d_str(zone.get("id"))].get(row_idx, [])
            cell.text = "\n".join(_node_cell_text(n) for n in nodes)
    return _Drawn(zones=len(zones), node_ids=placed, rows_clamped=rows_clamped)


def _add_diagram_walkthrough(doc, model: dict) -> None:
    """The phase-by-phase narrative the animation steps through — the part of an attack path that a
    still frame genuinely cannot carry, and the part a report reader most needs."""
    phases = [p for p in _d_list(model.get("phases")) if isinstance(p, dict)]
    phases.sort(key=lambda p: _d_int(p.get("n")))
    phases = phases[:_MAX_DIAGRAM_PHASES]
    if not phases:
        return
    doc.add_paragraph("Walkthrough", style="Heading 3")
    for phase in phases:
        title = _d_str(phase.get("title")) or ("Overview" if phase.get("intro") else "")
        head = f"Phase {_d_int(phase.get('n'))}"
        if title:
            head += f" — {title}"
        mitre = _d_str(phase.get("mitre"), 120)
        if mitre:
            head += f"  ({mitre})"
        p = doc.add_paragraph()
        p.add_run(head).bold = True
        desc = _d_str(phase.get("desc"), 1200)
        if desc:
            doc.add_paragraph(desc)


def _add_diagram_edges(doc, model: dict, drawn: _Drawn | None) -> None:
    """The connections, resolved to node LABELS — an edge list of bare ids is not a deliverable.

    Scoped to the nodes the GRID drew. Resolving against every node in the model let a connection name
    a host that appears nowhere in the table (its zone fell past the zone cap), which is a worse
    deliverable than omitting the line."""
    if drawn is None:
        return
    labels = {
        _d_str(n.get("id")): (_d_str(n.get("label")) or _d_str(n.get("id")))
        for n in _d_list(model.get("nodes"))[:_MAX_DIAGRAM_NODES]
        if isinstance(n, dict) and _d_str(n.get("id")) in drawn.node_ids
    }
    rendered = []
    # Sliced AFTER the dangling filter, not before: 80 leading dangling edges used to consume the whole
    # budget and print an empty Connections list while real edges sat behind them.
    for edge in _d_list(model.get("edges")):
        if len(rendered) >= _MAX_DIAGRAM_EDGES:
            break
        if not isinstance(edge, dict):
            continue
        src, dst = _d_str(edge.get("from")), _d_str(edge.get("to"))
        if src not in labels or dst not in labels:
            continue  # vector drops dangling edges too — do not draw a line to nowhere
        text = f"{labels[src]} → {labels[dst]}"
        detail = " · ".join(x for x in (_d_str(edge.get("label")), _d_str(edge.get("kind"), 60)) if x)
        rendered.append(f"{text}  ({detail})" if detail else text)
    if not rendered:
        return
    doc.add_paragraph("Connections", style="Heading 3")
    for line in rendered:
        doc.add_paragraph(line, style="List Bullet")


def _append_attack_paths(doc, ctx: ReportContext) -> None:
    """Append the Attack Paths section to the RENDERED document (ext#115), mirroring
    ``_append_checklists``/``_append_evidence_appendix``: programmatic and post-render, no Jinja loop
    authored into the binary template.

    Placed BEFORE the checklists and the evidence appendix so the .docx section order matches the HTML
    templates' (``findings`` -> ``diagrams`` -> ``methodology`` -> ``evidence``, see
    ``reporting/layouts.py``) — which is also what makes ``context.number_figures``'s single
    numbering sequence correct for both deliverables.

    Renders nothing when the engagement has no linked diagram, so a report without one is byte-identical
    to before this section existed (the same backward-compat guarantee ``render_html._render_diagrams``
    makes on the HTML side)."""
    if not ctx.diagrams:
        return
    doc.add_heading("Attack Paths", level=1)
    doc.add_paragraph(
        "Static rendition of each interactive attack path delivered with the HTML report: the zones "
        "and hosts as placed in the diagram, and the phase walkthrough, at the final step."
    )
    # Nothing caps how many diagrams may be linked to an engagement (``scribble_link_attack_path``
    # appends unconditionally; since #133 the operator's non-destructive out is ``include_in_report``,
    # which ``ReportContext`` already honours), so cap what one section renders -- same reasoning as
    # render_html's ``_MAX_APPENDIX_ITEMS``, and SAY SO in the document rather than truncating silently.
    withheld = max(0, len(ctx.diagrams) - _MAX_DIAGRAMS)
    if withheld:
        note = doc.add_paragraph()
        note.add_run(
            f"{withheld} further diagram{'s' if withheld != 1 else ''} linked to this engagement "
            f"{'are' if withheld != 1 else 'is'} not shown here "
            f"(this section lists at most {_MAX_DIAGRAMS})."
        ).italic = True
    budget = [_MAX_REPORT_SCAN_CHARS]
    for index, d in enumerate(ctx.diagrams[:_MAX_DIAGRAMS], start=1):
        model = _diagram_model(d.embed_html, budget)
        # ``meta`` was only guarded against None: a snapshot carrying ``"meta": "oops"`` (or a list)
        # reached ``.get`` on a str and raised AttributeError -- an uncaught 500 on EVERY future .docx
        # export of that engagement, from one stored PAT write, while the HTML export kept working.
        # That is this branch's own defect class inverted, so it is a dict or it is nothing.
        raw_meta = model.get("meta") if isinstance(model, dict) else None
        meta = raw_meta if isinstance(raw_meta, dict) else {}
        # ``d.caption`` comes from one of the two write paths (``POST …/attack-paths`` or, since #133,
        # ``PATCH …/attack-paths/<id>``), both of which strip NUL but not the other XML-illegal
        # control characters -- and those reach lxml unfiltered through add_heading/add_run.
        # The HEADING may use the model's own title -- more informative than a bare fallback. The
        # CAPTION may not: ``render_html`` falls back to the literal "Attack path", so using meta.title
        # here would print one figure number under two different captions across the two deliverables,
        # which is exactly what ext#117 says must not happen.
        title = _xml_safe(d.caption or "") or _d_str(meta.get("title")) or f"Attack path {index}"
        caption_text = _xml_safe(d.caption or "") or DIAGRAM_CAPTION_FALLBACK
        doc.add_heading(title, level=2)
        subtitle = _d_str(meta.get("subtitle"))
        if subtitle:
            doc.add_paragraph(subtitle)
        drawn = _add_diagram_grid(doc, model) if model else None
        if drawn is None:
            # Be loud about a snapshot this renderer could not read: an honest "the diagram is over
            # there" beats the silent omission ext#115 is about. It goes HERE, where the figure would
            # have been, and names ``diagram_ref`` -- for a reader holding only the .docx that is the
            # sole handle on which diagram this was.
            ref = _d_str(d.diagram_ref, 80)
            note = doc.add_paragraph()
            note.add_run(
                "This diagram is delivered as an interactive figure in the HTML report; a static "
                "rendition of its layout is not available here."
                + (f" (diagram {ref})" if ref else "")
            ).italic = True
        if model:
            _add_diagram_walkthrough(doc, model)
            _add_diagram_edges(doc, model, drawn)
            _note_truncation(doc, model, drawn)
        # The caption goes BENEATH the figure, carrying the same continuous number AND the same text
        # the HTML prints. "Caption" is a real Word style, so the hand-off can build a List of Figures.
        caption = _styled_paragraph(doc, "Caption")
        caption.add_run(figure_caption(d.figure_number, caption_text)).italic = True


def _append_checklists(doc, ctx: ReportContext) -> None:
    """Append the checklist sections to the RENDERED document with python-docx, rather than authoring a
    Jinja loop into the binary ``.docx`` template. Coverage/reminder -> a "Methodology and Coverage"
    section (status + note per item, grouped by section); compliance -> a "Compliance Attestation"
    section with a per-framework table (Control / Requirement / Result / Notes). Only checklists that
    opted into the report reach ``ctx.checklists``."""
    if not ctx.checklists:
        return
    coverage = [c for c in ctx.checklists if c.kind != "compliance"]
    compliance = [c for c in ctx.checklists if c.kind == "compliance"]

    def _rollup_line(cl) -> str:
        labels = [("satisfied", "Satisfied"), ("deficient", "Deficient"),
                  ("not_applicable", "N/A"), ("open", "Open")]
        return "   ".join(f"{lab}: {cl.rollup.get(k, 0)}" for k, lab in labels)

    if coverage:
        doc.add_heading("Methodology and Coverage", level=1)
        for cl in coverage:
            doc.add_heading(cl.name, level=2)
            doc.add_paragraph(_rollup_line(cl))
            last_section = object()
            for it in cl.items:
                if it.section != last_section:
                    last_section = it.section
                    if it.section:
                        doc.add_heading(it.section, level=3)
                p = doc.add_paragraph()
                p.add_run(f"[{(it.status or it.bucket_label)}] ").bold = True
                p.add_run(it.text)
                if it.finding_id and it.finding_title:
                    p.add_run(f"  (see: {it.finding_title})").italic = True
                if it.note:
                    doc.add_paragraph(it.note)

    if compliance:
        doc.add_heading("Compliance Attestation", level=1)
        for cl in compliance:
            doc.add_heading(cl.name, level=2)
            doc.add_paragraph(_rollup_line(cl))
            by_fw: dict[str, list] = {}
            fw_order: list[str] = []
            for it in cl.items:
                fw = it.framework or ""
                if fw not in by_fw:
                    by_fw[fw] = []
                    fw_order.append(fw)
                by_fw[fw].append(it)
            for fw in fw_order:
                if fw:
                    doc.add_heading(fw, level=3)
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
                    "Control", "Requirement", "Result", "Notes")
                for it in by_fw[fw]:
                    cells = table.add_row().cells
                    cells[0].text = it.control_ref or ""
                    cells[1].text = it.text
                    cells[2].text = it.bucket_label
                    note = it.note or ""
                    if it.finding_id and it.finding_title:
                        note = (note + f"  (see: {it.finding_title})").strip()
                    cells[3].text = note


def _append_evidence_appendix(doc, ctx: ReportContext, *, artifact_bytes: ArtifactBytes | None) -> None:
    """Append engagement-level evidence (``ReportContext.artifacts`` -- artifacts attached to the
    ENGAGEMENT with no ``finding_id``) to the RENDERED document, mirroring ``_append_checklists``
    (programmatic, post-render -- no Jinja loop authored into the binary template). This is the docx
    half of issue #54 / ext#40 mechanism 1: before this, such an upload was accepted, stored, answered
    201 with a URL, and then appeared in no docx deliverable at all -- the exact HTML-side gap
    ``_render_evidence_appendix`` (render_html.py) already closed.

    Only added when ``ctx.artifacts`` is non-empty, so an engagement with no engagement-level evidence
    renders exactly as it did before this section existed. Each item: an image within the size bound
    embeds via ``doc.add_picture``; anything else (non-image, oversized, or unavailable bytes) gets a
    caption/filename paragraph instead of the render failing. There is no zip delivery mode for docx
    (issue #62 does not apply here), so this is the one size bound this format needs."""
    if not ctx.artifacts:
        return
    doc.add_heading("Evidence Appendix", level=1)
    doc.add_paragraph(
        "Evidence recorded against this engagement as a whole rather than against one finding."
    )
    for a in ctx.artifacts:
        is_image = (a.content_type or "").startswith("image/")
        data: bytes | None = None
        if is_image and artifact_bytes:
            try:
                data = artifact_bytes(a.storage_path)
            except Exception:
                data = None
            if data and len(data) > _MAX_EVIDENCE_BYTES:
                data = None  # oversized: fall back to caption-only rather than embed a huge blob
        caption = _numbered_caption(a)  # ext#117
        if data:
            try:
                doc.add_picture(io.BytesIO(data), width=_EVIDENCE_IMAGE_WIDTH)
                doc.add_paragraph(caption)
                continue
            except Exception:
                pass  # fall through to the caption-only path below
        p = doc.add_paragraph()
        p.add_run(f"{caption} ").italic = True
        p.add_run(f"({a.filename} -- not embedded)").italic = True


def render_report_docx(ctx: ReportContext, *, artifact_bytes: ArtifactBytes | None = None) -> bytes:
    """Render ``ctx`` to a ``.docx`` document (bytes) using ``report_templates/default.docx``.

    ``artifact_bytes(storage_path) -> bytes | None`` supplies evidence-gallery + inline-content image
    bytes; when ``None`` (or a lookup fails), images degrade to caption-only / bracketed-placeholder
    text rather than the render failing.
    """
    tpl = DocxTemplate(str(_TEMPLATE_PATH))
    tpl.init_docx()
    # Content-block conversion happens eagerly (before ``tpl.render()``), so inline images need
    # ``current_rendering_part`` set now — this is the same part ``render()`` would assign for the
    # main document body, so it's consistent with what happens automatically for everything else
    # (scalars, evidence ``InlineImage``s) during the subsequent ``tpl.render()`` call.
    tpl.current_rendering_part = tpl.docx.part

    context = _build_context(ctx, tpl=tpl, artifact_bytes=artifact_bytes)
    tpl.render(context)
    # Section order mirrors the HTML templates' block order (findings -> diagrams -> methodology ->
    # evidence, reporting/layouts.py), which is what makes context.number_figures' single figure
    # sequence come out the same in both deliverables.
    _append_attack_paths(tpl.docx, ctx)  # ext#115
    _append_checklists(tpl.docx, ctx)  # programmatic, post-render (no Jinja in the binary template)
    _append_evidence_appendix(tpl.docx, ctx, artifact_bytes=artifact_bytes)

    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()
