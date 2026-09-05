"""ext#115 (BUG-7) — the attack path must not be SILENTLY DROPPED from the ``.docx`` deliverable.

Reported from the TeamsPlus engagement, retest 2026-08-26: the HTML export embedded the diagram as an
animated ``<iframe class="attack-path-frame">`` and played it, while the ``.docx`` rendered from the
SAME engagement contained no attack path at all — the string ``Attack path`` did not appear in
``word/document.xml`` and all 8 ``<w:drawing>`` elements were evidence screenshots. A reader given only
the Word deliverable had no way to know a diagram existed.

Why these assertions look the way they do
-----------------------------------------
This is a "renders fine, content missing" defect, so a test that only proves the render SUCCEEDED is
worthless — that test was already green throughout the bug. Every assertion below therefore names
CONTENT that must be present in the rendered document (the zone titles, the host labels, the phase
walkthrough, the caption), read back out of the real ``.docx`` rather than off the context.

``tests/test_report_attack_path.py`` covers the HTML side and the backward-compat guarantee; this file
is its ``.docx`` half plus the untrusted-snapshot bounds.
"""

from __future__ import annotations

import io
import json
import zipfile

import docx
from docx.oxml.ns import qn

from scribble.content import schema
from scribble.enums import Severity
from scribble.models import (
    AttackChain,
    AttackChainStep,
    Client,
    Engagement,
    EngagementDiagram,
    EngagementFinding,
    FindingGroup,
)
from scribble.reporting import build_report_context
from scribble.reporting.render_docx import render_report_docx

MODEL = {
    "schema": "vector.attackpath/v1",
    "meta": {"title": "Acme compromise chain", "subtitle": "external to domain admin"},
    "zones": [
        {"id": "ext", "title": "Internet", "order": 0},
        {"id": "dmz", "title": "DMZ", "order": 1},
        {"id": "core", "title": "Core AD", "order": 2},
    ],
    "nodes": [
        {"id": "kali", "label": "Operator", "zone": "ext", "row": 0, "ip": "203.0.113.9"},
        {"id": "web", "label": "www01", "zone": "dmz", "row": 0, "ip": "198.51.100.5",
         "states": [{"at": 1, "state": "owned"}]},
        {"id": "dc", "label": "DC01", "zone": "core", "row": 1, "domain": "acme.local",
         "states": [{"at": 1, "state": "target"}, {"at": 3, "state": "owned"}]},
    ],
    "edges": [
        {"id": "e1", "from": "kali", "to": "web", "kind": "attack", "at": 1, "label": "Struts RCE"},
        {"id": "e2", "from": "web", "to": "dc", "kind": "attack", "at": 3, "label": "Kerberoast"},
        # dangling: vector drops these and so must this renderer — never draw a line to nowhere.
        {"id": "e3", "from": "web", "to": "ghost", "kind": "attack", "at": 3},
    ],
    "phases": [
        {"n": 0, "intro": True, "title": "Setup"},
        {"n": 1, "title": "Initial access", "mitre": "T1190", "desc": "Exploit the edge web app."},
        {"n": 3, "title": "Domain compromise", "mitre": "T1558.003", "desc": "Kerberoast to DA."},
    ],
}


def _script_json(obj) -> str:
    """vector's ``render.json_for_script`` escaping: ``<``/``>``/``&`` become ``\\uXXXX``, which is what
    makes "match to the first ``</script>``" exact rather than best-effort."""
    esc = {ord("<"): "\\u003c", ord(">"): "\\u003e", ord("&"): "\\u0026"}
    return json.dumps(obj, ensure_ascii=False, separators=(",", ":")).translate(esc)


def snapshot(model=None) -> str:
    """A stand-in for vector's ``export.html`` — the same shape ``templates/vector/deliverable.html.j2``
    emits, including the ``<script type="application/json" id="vap-model">`` block this renderer reads."""
    return (
        "<!doctype html><html><head><title>Acme compromise chain</title></head><body>"
        '<div id="vap"></div>'
        f'<script type="application/json" id="vap-model">{_script_json(model or MODEL)}</script>'
        "<script>/* viewer runtime */</script></body></html>"
    )


def _block(text: str) -> dict:
    return schema.doc_from_text(text)


def _engagement(session_factory, *, diagrams: list[tuple[str, str]] | None = None):
    """``diagrams`` is a list of ``(caption, embed_html)``; ``None`` links none."""
    with session_factory() as db:
        client = Client(name="Acme Co")
        db.add(client)
        db.flush()
        eng = Engagement(name="TeamsPlus Assessment", client_id=client.id, company_name="Acme Corp")
        grp = FindingGroup(engagement=eng, name="External", order_index=0)
        EngagementFinding(
            engagement=eng, group=grp, title="Reflected XSS", severity=Severity.high,
            order_index=0, content_json={"description": _block("Reflected XSS on /search.")},
        )
        db.add(eng)
        db.flush()
        for i, (caption, embed) in enumerate(diagrams or []):
            db.add(EngagementDiagram(
                engagement_id=eng.id, diagram_ref=f"ref-{i}", caption=caption,
                embed_html=embed, order_index=i,
            ))
        db.commit()
        return eng.id


def _render(session_factory, eng_id) -> bytes:
    with session_factory() as db:
        eng = db.get(Engagement, eng_id)
        return render_report_docx(build_report_context(eng))


def _text(raw: bytes) -> str:
    doc = docx.Document(io.BytesIO(raw))
    return "".join(t.text or "" for t in doc.element.body.iter(qn("w:t")))


def _document_xml(raw: bytes) -> str:
    return zipfile.ZipFile(io.BytesIO(raw)).read("word/document.xml").decode("utf-8")


# ── the reported defect, asserted as the reporter measured it ────────────────────────────────────────


def test_docx_carries_the_attack_path_the_html_has(session_factory):
    """The headline of ext#115, phrased exactly as the bug report measured it: `grep 'Attack path'`
    over ``word/document.xml``."""
    eng_id = _engagement(session_factory, diagrams=[("Domain compromise chain", snapshot())])
    raw = _render(session_factory, eng_id)
    xml = _document_xml(raw)
    assert "Attack Paths" in xml
    assert "attack path" in xml.lower()


def test_docx_carries_the_diagram_caption_beneath_the_figure(session_factory):
    """"…with the caption beneath it" (ext#115). Order matters: the caption is a CAPTION, not a title,
    so it must follow the rendition rather than lead it."""
    eng_id = _engagement(session_factory, diagrams=[("Domain compromise chain", snapshot())])
    doc = docx.Document(io.BytesIO(_render(session_factory, eng_id)))
    texts = [p.text for p in doc.paragraphs]
    heading = texts.index("Domain compromise chain")
    caption = next(i for i, t in enumerate(texts) if t.startswith("Figure ") and "Domain" in t)
    assert caption > heading
    # …and after the walkthrough, i.e. beneath the whole figure rather than tucked under its title.
    assert caption > texts.index("Walkthrough")


def test_docx_carries_the_diagram_topology_and_walkthrough(session_factory):
    """The static frame itself: zones as columns, hosts placed in them, the phase narrative, and the
    connections resolved to LABELS. Each of these is content the Word reader silently lost before."""
    eng_id = _engagement(session_factory, diagrams=[("Domain compromise chain", snapshot())])
    text = _text(_render(session_factory, eng_id))
    for zone in ("Internet", "DMZ", "Core AD"):
        assert zone in text, f"zone {zone!r} missing from the docx"
    for host in ("Operator", "www01", "DC01"):
        assert host in text, f"host {host!r} missing from the docx"
    assert "203.0.113.9" in text and "acme.local" in text
    assert "[OWNED]" in text  # the FINAL state, not the first: a still frame shows the end of the path
    assert "Phase 1 — Initial access" in text and "T1190" in text
    assert "Exploit the edge web app." in text
    assert "Operator → www01" in text and "Struts RCE" in text
    assert "ghost" not in text  # the dangling edge is dropped, not drawn to nowhere


def test_docx_names_the_subtitle_and_falls_back_to_the_model_title(session_factory):
    """An operator who linked a diagram without typing a caption still gets a titled figure — the
    model's own ``meta.title`` — rather than an anonymous table."""
    eng_id = _engagement(session_factory, diagrams=[("", snapshot())])
    text = _text(_render(session_factory, eng_id))
    assert "Acme compromise chain" in text
    assert "external to domain admin" in text


# ── backward compatibility: an engagement with no diagram is untouched ───────────────────────────────


def test_no_diagram_appends_no_attack_paths_section(session_factory):
    """The same guarantee ``render_html._render_diagrams`` makes on the HTML side. Without it, every
    existing report would grow an empty section."""
    eng_id = _engagement(session_factory, diagrams=None)
    xml = _document_xml(_render(session_factory, eng_id))
    assert "Attack Paths" not in xml
    assert "Walkthrough" not in xml


# ── an unreadable snapshot degrades LOUDLY, never silently ───────────────────────────────────────────


def test_an_unreadable_snapshot_still_announces_the_diagram(session_factory):
    """``embed_html`` is stored verbatim from a PAT POST, so nothing guarantees it came from vector.
    When the model cannot be read the section must still say the diagram EXISTS — reverting to silence
    is the whole defect."""
    eng_id = _engagement(session_factory, diagrams=[("Chain of pwnage", "<html>not vector</html>")])
    text = _text(_render(session_factory, eng_id))
    assert "Attack Paths" in text
    assert "Chain of pwnage" in text
    assert "interactive figure in the HTML report" in text


def test_malformed_json_in_the_model_block_degrades_gracefully(session_factory):
    bad = '<html><body><script type="application/json" id="vap-model">{not json</script></body></html>'
    eng_id = _engagement(session_factory, diagrams=[("Broken", bad)])
    text = _text(_render(session_factory, eng_id))
    assert "Broken" in text and "interactive figure in the HTML report" in text


def test_a_model_with_no_zones_still_renders_its_walkthrough(session_factory):
    """Nothing placeable to draw, but the phases are still the most useful half of an attack path."""
    model = {"meta": {"title": "Narrative only"}, "zones": [], "nodes": [], "edges": [],
             "phases": [{"n": 1, "title": "Initial access", "desc": "Phished a user."}]}
    eng_id = _engagement(session_factory, diagrams=[("Narrative", snapshot(model))])
    text = _text(_render(session_factory, eng_id))
    assert "Phase 1 — Initial access" in text and "Phished a user." in text


# ── bounds on an untrusted snapshot ──────────────────────────────────────────────────────────────────


def test_a_hostile_snapshot_cannot_render_an_unbounded_table(session_factory):
    """``embed_html`` is operator/agent-supplied and vector's own caps do not apply to it (nothing
    proves it came from vector). Cap zones/rows here or one POST turns every report render into a
    thousand-column table."""
    from scribble.reporting.render_docx import _MAX_DIAGRAM_ZONES

    model = {
        "meta": {"title": "Flood"},
        "zones": [{"id": f"z{i}", "title": f"Zone {i}", "order": i} for i in range(400)],
        "nodes": [{"id": f"n{i}", "zone": f"z{i}", "label": f"host{i}", "row": i} for i in range(400)],
        "edges": [], "phases": [],
    }
    eng_id = _engagement(session_factory, diagrams=[("Flood", snapshot(model))])
    doc = docx.Document(io.BytesIO(_render(session_factory, eng_id)))
    assert doc.tables, "expected the zone grid to render at all"
    grid = doc.tables[-1]
    assert len(grid.columns) == _MAX_DIAGRAM_ZONES
    assert len(grid.rows) <= 41  # header + _MAX_DIAGRAM_ROWS


def test_an_oversized_snapshot_is_not_scanned(session_factory):
    """A snapshot past the scan bound is not walked at all — it degrades to the honest note rather
    than costing a multi-megabyte scan on every report render."""
    from scribble.reporting.render_docx import _MAX_DIAGRAM_SCAN_CHARS

    huge = "<html>" + ("x" * (_MAX_DIAGRAM_SCAN_CHARS + 1)) + "</html>"
    eng_id = _engagement(session_factory, diagrams=[("Huge", huge)])
    text = _text(_render(session_factory, eng_id))
    assert "Huge" in text and "interactive figure in the HTML report" in text


def test_untrusted_model_fields_are_coerced_not_trusted(session_factory):
    """Every field is coerced/capped on the way out: a dict where a label belongs, a 10k-character
    description, a bool where a number belongs. None of these may reach python-docx as-is."""
    model = {
        "meta": {"title": {"nope": 1}, "subtitle": ["also nope"]},
        "zones": [{"id": "z", "title": "Z" * 5000, "order": "not-an-int"}],
        "nodes": [{"id": "n", "zone": "z", "label": {"x": 1}, "row": True, "states": "not-a-list"}],
        "edges": [{"id": "e", "from": "n", "to": "n", "label": 12345}],
        "phases": [{"n": "x", "title": None, "desc": "d" * 9000}],
    }
    eng_id = _engagement(session_factory, diagrams=[("Coerced", snapshot(model))])
    text = _text(_render(session_factory, eng_id))
    assert "Coerced" in text
    assert "Z" * 401 not in text  # the 5000-char zone title was capped
    assert "d" * 1201 not in text


# ── hostile-snapshot hardening (found by the branch's independent security review) ───────────────────

# XML-illegal control characters, written as ESCAPES so this file itself stays clean text.
_CTRL = "\x00\x01\x02\x07\x08\x0b\x1b\x7f"


def test_a_pathological_snapshot_does_not_wedge_the_renderer(session_factory):
    """ReDoS. Extraction used to be a regex with two unanchored ``[^>]*`` runs before a literal, so
    every ``<script`` in the input was a start position costing O(n) -- measured 15.9s at 64 KiB of
    ``"<script " * 8000``, extrapolating to ~71 minutes at 1 MiB, which the link route's 10 MiB cap
    happily accepts. CPython's ``re`` holds the GIL and never yields to the gevent hub, so that is not
    a slow request, it is the whole worker; and lowering the size cap does not fix a quadratic.

    A wall-clock budget is the only assertion that can distinguish "linear" from "quadratic" here. It
    is generous (5s for work that measures in milliseconds) precisely so it fails on the ALGORITHM and
    not on a loaded box."""
    import time

    hostile = "<script " * 8000  # the exact shape that measured 15.9s against the old pattern
    eng_id = _engagement(session_factory, diagrams=[("Pathological", hostile)])
    started = time.monotonic()
    text = _text(_render(session_factory, eng_id))
    elapsed = time.monotonic() - started
    assert elapsed < 5.0, f"extraction took {elapsed:.1f}s -- the scan is not linear"
    assert "Pathological" in text and "interactive figure in the HTML report" in text


def test_json_float_specials_do_not_500_the_deliverable(session_factory):
    """``json`` accepts the non-standard ``Infinity`` and overflows ``1e999`` to ``float("inf")``;
    ``int(float("inf"))`` raises ``OverflowError``, which is an ``ArithmeticError`` and was NOT in the
    coercion's ``except``. One token in a linked snapshot made every future ``.docx`` export of that
    engagement an uncaught 500 -- defeating the section's own "degrade, never raise" contract."""
    model = {
        "meta": {"title": "Infinities"},
        "zones": [{"id": "z", "title": "Zed", "order": 1e999}],
        "nodes": [{"id": "n", "zone": "z", "label": "host", "row": 1e999,
                   "states": [{"at": 1e999, "state": "owned"}]}],
        "edges": [], "phases": [{"n": 1e999, "title": "Phase", "desc": "dee"}],
    }
    eng_id = _engagement(session_factory, diagrams=[("Infinities", snapshot(model))])
    text = _text(_render(session_factory, eng_id))  # must not raise
    assert "Infinities" in text and "Zed" in text and "host" in text


def test_xml_illegal_control_characters_do_not_500_the_deliverable(session_factory):
    """lxml (under python-docx) refuses 23 codepoints that are legal in a Python str and legal in
    JSON. They cannot be filtered at the link route: the JSON blob carries them as the six ASCII
    characters ``\\u0000``, so the stored ``embed_html`` holds no literal control byte for the route's
    NUL scrub to find, and ``json.loads`` materializes the real character here. Both the model fields
    and the operator-typed CAPTION are affected -- the caption never goes through the model coercion
    at all."""
    model = {
        "meta": {"title": "Ctrl" + _CTRL},
        "zones": [{"id": "z", "title": "Zone" + _CTRL, "order": 0}],
        "nodes": [{"id": "n", "zone": "z", "label": "host" + _CTRL, "ip": "10.0.0.1"}],
        "edges": [], "phases": [{"n": 1, "title": "Phase" + _CTRL, "desc": "desc" + _CTRL}],
    }
    eng_id = _engagement(
        session_factory, diagrams=[("Caption" + _CTRL, snapshot(model))]
    )
    text = _text(_render(session_factory, eng_id))  # must not raise
    assert "Caption" in text
    assert "Zone" in text and "host" in text and "Phase" in text
    for bad in _CTRL:
        assert bad not in text


def test_the_node_count_is_capped(session_factory):
    """``nodes`` was the one list with no cap, and every node sharing a ``(zone, row)`` concatenates
    into ONE cell -- one ``<w:br/>`` element each. Measured uncapped: 100k nodes = 8.8s and 204 MB
    peak per render, multiplied by however many diagrams are linked."""
    from scribble.reporting.render_docx import _MAX_DIAGRAM_NODES

    model = {
        "meta": {"title": "Node flood"},
        "zones": [{"id": "z", "title": "Zone", "order": 0}],
        "nodes": [{"id": f"n{i}", "zone": "z", "label": f"host{i}", "row": 0}
                  for i in range(_MAX_DIAGRAM_NODES + 500)],
        "edges": [], "phases": [],
    }
    eng_id = _engagement(session_factory, diagrams=[("Flood", snapshot(model))])
    text = _text(_render(session_factory, eng_id))
    assert f"host{_MAX_DIAGRAM_NODES - 1}" in text
    assert f"host{_MAX_DIAGRAM_NODES}" not in text


def test_the_diagram_count_is_capped_and_says_so(session_factory):
    """Nothing caps how many diagrams may be linked to one engagement, so cap what one section
    renders -- and NAME the shortfall, because a silently truncated section is the same defect class
    this whole branch exists to fix."""
    from scribble.reporting.render_docx import _MAX_DIAGRAMS

    eng_id = _engagement(
        session_factory,
        diagrams=[(f"Diagram {i}", snapshot()) for i in range(_MAX_DIAGRAMS + 3)],
    )
    text = _text(_render(session_factory, eng_id))
    assert f"Diagram {_MAX_DIAGRAMS - 1}" in text
    assert f"Diagram {_MAX_DIAGRAMS + 2}" not in text
    assert "3 further diagrams" in text


def test_a_non_dict_meta_does_not_500_the_deliverable(session_factory):
    """``meta`` was guarded against ``None`` only, so a snapshot carrying ``"meta": "oops"`` reached
    ``.get`` on a ``str`` and raised ``AttributeError`` -- an uncaught 500 on every future ``.docx``
    export of that engagement from one stored PAT write, while the HTML export kept working. Every
    other JSON shape a hostile snapshot can put there is covered here too, in ONE engagement so a
    single render has to survive all of them."""
    def _m(bad_meta):
        return {"meta": bad_meta,
                "zones": [{"id": "z", "title": "Zed", "order": 0}],
                "nodes": [{"id": "n", "zone": "z", "label": "host"}],
                "edges": [], "phases": []}

    eng_id = _engagement(session_factory, diagrams=[
        (f"Meta {i}", snapshot(_m(bad)))
        for i, bad in enumerate(("oops", ["a", "list"], 7, True, None))
    ])
    text = _text(_render(session_factory, eng_id))  # must not raise
    assert text.count("Zed") == 5, "one of the five hostile `meta` shapes lost its diagram"


def test_xml_illegal_surrogates_and_noncharacters_do_not_500_the_deliverable(session_factory):
    """The first hardening pass derived its character class from C0 alone. Two more classes survive
    ``json.loads`` and still kill lxml: lone surrogates ``D800``-``DFFF`` (``UnicodeEncodeError`` out
    of lxml's utf-8 encode) and the noncharacters ``FFFE``/``FFFF`` (``ValueError``). Same permanent
    500 on every future ``.docx`` export, from one stored write.

    Note how they must arrive: a lone surrogate is not encodable, so it can NEVER be a literal
    character in a stored ``embed_html`` -- it comes in as the six ASCII characters of a JSON escape,
    which is exactly why the route's scrub cannot see it. So this snapshot is built ASCII-only."""
    bad = "\ud800\udfff\ufffe\uffff"
    model = {
        "meta": {"title": "Surrogates" + bad, "subtitle": "sub" + bad},
        "zones": [{"id": "z", "title": "Zone" + bad, "order": 0}],
        "nodes": [{"id": "n", "zone": "z", "label": "host" + bad,
                   "states": [{"at": 1, "label": "state" + bad}]}],
        "edges": [], "phases": [{"n": 1, "title": "Phase" + bad, "desc": "desc" + bad}],
    }
    embed = (
        "<!doctype html><html><body>"
        f'<script type="application/json" id="vap-model">'
        f"{json.dumps(model, ensure_ascii=True, separators=(',', ':'))}</script></body></html>"
    )
    assert embed.isascii(), "the stored snapshot must be pure ASCII -- that is the whole point"
    eng_id = _engagement(session_factory, diagrams=[("Surrogates", embed)])
    text = _text(_render(session_factory, eng_id))  # must not raise
    assert "Zone" in text and "host" in text and "Phase" in text
    for ch in bad:
        assert ch not in text


def test_the_int_coercion_survives_an_infinite_float():
    """The parser kills ``Infinity`` before it can reach here, so the ``OverflowError`` arm of the
    coercion is unreachable through a snapshot -- which means the end-to-end test above cannot prove
    it. Asserted directly instead, because ``int(float("inf"))`` raises ``OverflowError``, an
    ``ArithmeticError`` and NOT a ``ValueError``: the day a second caller feeds ``_d_int`` from
    somewhere other than ``parse_constant``'d JSON, the missing arm is a 500 again."""
    from scribble.reporting.render_docx import _d_int

    assert _d_int(float("inf")) == 0
    assert _d_int(float("-inf"), default=7) == 7
    assert _d_int(float("nan")) == 0


def test_the_report_wide_scan_budget_bounds_the_whole_section(session_factory):
    """``_MAX_DIAGRAM_SCAN_CHARS`` bounds ONE snapshot; nothing bounded the product. 50 linked
    diagrams of 9.6 MiB each measured 34s of GIL-held CPU and 515 MiB peak per ``.docx`` GET -- rows
    any ``write``-scope PAT can store and any ``read``-scope viewer can then trigger, repeatedly.

    The budget is spent across the section, so a diagram arriving after it is exhausted degrades to
    the same honest note a snapshot the renderer cannot read already gets."""
    from scribble.reporting.render_docx import _MAX_REPORT_SCAN_CHARS

    filler = "<!--" + ("x" * (_MAX_REPORT_SCAN_CHARS // 2)) + "-->"
    eng_id = _engagement(session_factory, diagrams=[
        ("First", filler + snapshot()),
        ("Second", filler + snapshot()),
        ("Third", filler + snapshot()),
    ])
    text = _text(_render(session_factory, eng_id))
    # All three are still ANNOUNCED -- the section never goes silent, which is the whole point of #115.
    assert "First" in text and "Second" in text and "Third" in text
    # ...but the third was past the budget, so it degrades to the note instead of being scanned.
    assert "interactive figure in the HTML report" in text


def test_a_decoy_model_id_cannot_hide_the_real_diagram(session_factory):
    """Extraction took the FIRST ``id="vap-model"`` only and gave up if it was not inside a
    ``<script>``. A snapshot that puts a decoy ahead of the genuine block therefore made Word print
    "static rendition not available" while the HTML iframe drew the diagram in full -- the Word reader
    silently getting less, which is ext#115 restated."""
    decoy = '<!-- id="vap-model" --><div id="vap-model">not json</div>'
    eng_id = _engagement(session_factory, diagrams=[("Decoy", decoy + snapshot())])
    text = _text(_render(session_factory, eng_id))
    assert "Internet" in text and "DC01" in text, "the decoy hid the real model"
    assert "interactive figure in the HTML report" not in text


def test_the_truncation_note_counts_what_was_DRAWN_not_the_caps(session_factory):
    """The note used to count ``len(nodes) - _MAX_DIAGRAM_NODES``, which under-reports -- and it
    under-reports the HOSTS specifically. A node is also dropped when its zone fell past the zone cap,
    when its zone id is unknown, or when two zones shared an id and de-duped. On this very fixture the
    old arithmetic announced "388 zones" and said nothing at all about the 388 hosts that went with
    them; a silently truncated figure is the defect class this branch exists to fix."""
    from scribble.reporting.render_docx import _MAX_DIAGRAM_ZONES

    model = {
        "meta": {"title": "Flood"},
        "zones": [{"id": f"z{i}", "title": f"Zone {i}", "order": i} for i in range(400)],
        "nodes": [{"id": f"n{i}", "zone": f"z{i}", "label": f"host{i}", "row": 0} for i in range(400)],
        "edges": [], "phases": [],
    }
    eng_id = _engagement(session_factory, diagrams=[("Flood", snapshot(model))])
    text = _text(_render(session_factory, eng_id))
    assert f"{400 - _MAX_DIAGRAM_ZONES} zones" in text
    assert f"{400 - _MAX_DIAGRAM_ZONES} hosts" in text, "the vanished hosts were not named"


def test_a_connection_never_names_a_host_that_is_not_in_the_table(session_factory):
    """The edge list resolved labels from EVERY node in the model, so a connection could name a host
    the reader cannot find anywhere in the figure -- its zone having fallen past the zone cap. A
    dangling reference in a client deliverable is worse than an omitted line."""
    from scribble.reporting.render_docx import _MAX_DIAGRAM_ZONES

    model = {
        "meta": {"title": "Offscreen"},
        "zones": [{"id": f"z{i}", "title": f"Zone {i}", "order": i}
                  for i in range(_MAX_DIAGRAM_ZONES + 1)],
        "nodes": [{"id": "shown", "zone": "z0", "label": "shownhost", "row": 0},
                  {"id": "hidden", "zone": f"z{_MAX_DIAGRAM_ZONES}", "label": "hiddenhost", "row": 0}],
        "edges": [{"id": "e", "from": "shown", "to": "hidden", "kind": "attack"}],
        "phases": [],
    }
    eng_id = _engagement(session_factory, diagrams=[("Offscreen", snapshot(model))])
    text = _text(_render(session_factory, eng_id))
    assert "shownhost" in text
    assert "hiddenhost" not in text, "a connection named a host that is nowhere in the table"


def test_the_status_chip_matches_what_the_viewer_prints(session_factory):
    """The chip must read the way ``vector-viewer.js``'s ``nodeVisual()`` renders it at the final
    keyframe, because the section's claim is that the Word table is the same figure. Its rule is
    three-tiered and an earlier pass here inverted all three: an explicit ``states[].label`` wins
    outright and prints VERBATIM; otherwise the highest-precedence state that is actually IN the
    catalog prints the catalog's display label; an UNKNOWN state key is not a candidate at all (the
    viewer shows nothing for it); and with no state the node's ROLE supplies the chip."""
    model = {
        "meta": {"title": "Chips"},
        "zones": [{"id": "z", "title": "Zone", "order": 0}],
        "nodes": [
            # 1. explicit label beats the catalog, and is NOT uppercased
            {"id": "a", "zone": "z", "label": "alpha", "row": 0,
             "states": [{"at": 1, "state": "owned", "label": "Domain Admin"}]},
            # 2. catalog display label, highest precedence wins over a later lower one
            {"id": "b", "zone": "z", "label": "bravo", "row": 1,
             "states": [{"at": 1, "state": "impacted"}, {"at": 2, "state": "target"}]},
            # 3. an unknown state key is not a candidate -- no chip
            {"id": "c", "zone": "z", "label": "charlie", "row": 2,
             "states": [{"at": 1, "state": "pwned"}]},
            # 4. no state at all -> the role's status text
            {"id": "d", "zone": "z", "label": "delta", "row": 3, "role": "c2"},
        ],
        "edges": [], "phases": [],
    }
    eng_id = _engagement(session_factory, diagrams=[("Chips", snapshot(model))])
    text = _text(_render(session_factory, eng_id))
    assert "Domain Admin" in text and "[OWNED]" not in text  # 1
    assert "[IMPACT]" in text and "[TARGET]" not in text     # 2
    assert "[PWNED]" not in text                              # 3
    assert "[C2]" in text                                     # 4


# ── #628: the attack-chain NARRATIVE mirrors into the .docx ───────────────────────────────────────────


def _chain_engagement(session_factory, *, embed: bool = False) -> str:
    with session_factory() as db:
        client = Client(name="Acme Co")
        db.add(client)
        db.flush()
        eng = Engagement(name="Chain Assessment", client_id=client.id, company_name="Acme Corp")
        grp = FindingGroup(engagement=eng, name="External", order_index=0)
        EngagementFinding(
            engagement=eng, group=grp, title="Reflected XSS", severity=Severity.high,
            order_index=0, content_json={"description": _block("Reflected XSS on /search.")},
        )
        db.add(eng)
        db.flush()
        chain = AttackChain(
            engagement_id=eng.id, title="Foothold to Domain Admin",
            summary="Chained the edge XSS into a full domain compromise.",
            order_index=0, embed_html=(snapshot() if embed else None),
        )
        db.add(chain)
        db.flush()
        for oi, title in ((1, "Pivot to the DMZ web host"), (0, "Initial access via reflected XSS")):
            db.add(AttackChainStep(chain_id=chain.id, order_index=oi, title=title))
        db.commit()
        return eng.id


def test_docx_carries_the_attack_chain_narrative(session_factory):
    """A "renders fine, content missing" guard like the diagram one: the .docx must actually contain the
    chain's title, summary and every step title, read back out of ``word/document.xml`` — and the steps
    must be numbered in ``order_index`` order, not insertion order."""
    text = _text(_render(session_factory, _chain_engagement(session_factory)))
    assert "Attack Chains" in text
    assert "Foothold to Domain Admin" in text
    assert "Chained the edge XSS into a full domain compromise." in text
    assert "1. Initial access via reflected XSS" in text
    assert "2. Pivot to the DMZ web host" in text


def test_docx_notes_a_chain_embed_lives_in_the_html_report(session_factory):
    """A chain's OPTIONAL ``embed_html`` visual is HTML-only in the .docx; the Word deliverable carries an
    explicit pointer so nothing is SILENTLY dropped (the ext#115 defect class, applied to chains)."""
    text = _text(_render(session_factory, _chain_engagement(session_factory, embed=True)))
    assert "interactive figure in the HTML" in text


def test_docx_without_a_chain_has_no_attack_chains_section(session_factory):
    """Backward-compat: an engagement with no chain is byte-identical to before this section existed."""
    text = _text(_render(session_factory, _engagement(session_factory)))
    assert "Attack Chains" not in text
