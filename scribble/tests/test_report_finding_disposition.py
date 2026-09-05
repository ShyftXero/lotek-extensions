"""lotek#618: a finding's ``status`` reaches the report, and drives ONE report-disposition predicate.

The bug these tests pin is not a missing badge. Inclusion in a report was decided by
``include_in_report`` **alone**, and ``build_report_context._tally`` counted every included finding's
severity into ``SeverityRollup`` -> ``risk_rating()`` -> the risk banner and the generated narrative.
So a finding the operator had marked ``false_positive`` or ``fixed`` **inflated the client's overall
risk rating** — the report asserted a number that was not true.

``report_disposition()`` (``scribble.enums``) is the single home for that derivation; every surface
(context, HTML, DOCX) reads it rather than comparing ``status`` itself. The cross-product test below is
what proves the surfaces agree: a report that shows "Remediated" while the banner still counts the
finding as live is exactly the drift a second copy of the predicate produces.
"""

from __future__ import annotations

import io
import uuid

import docx
import pytest

from scribble.content import schema
from scribble.enums import (
    FindingStatus,
    Severity,
    finding_status_label,
    report_disposition,
)
from scribble.models import Client, Engagement, EngagementFinding, FindingGroup
from scribble.reporting import build_report_context
from scribble.reporting.render_docx import render_report_docx
from scribble.reporting.render_html import render_report_html

# status -> (disposition, client-facing label, reaches the deliverable?, counted in the ladder?)
#
# Deliberate label choices (lotek#618, Decision 3): "Remediated" NOT "Fixed (verified)", and
# "Risk accepted" NOT "Accepted risk (client decision)" — both of the rejected forms assert work
# nobody recorded (a verification; a client sign-off), which is what
# ``tests/test_report_standing_prose.py`` exists to forbid. The verification wording belongs to the
# retest model (lotek#621), where a date and a verifier actually exist.
EXPECTED = {
    FindingStatus.new: ("live", "", True, True),
    FindingStatus.triaged: ("live", "Triaged", True, True),
    FindingStatus.needs_retest: ("live", "Awaiting retest", True, True),
    FindingStatus.fixed: ("remediated", "Remediated", True, False),
    FindingStatus.accepted_risk: ("accepted", "Risk accepted", True, False),
    FindingStatus.false_positive: ("excluded", "", False, False),
}


def _block(text: str) -> dict:
    return schema.doc_from_text(text)


def _all_text(document) -> str:
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _engagement(
    session_factory, findings: list[tuple[str, Severity, FindingStatus, bool]]
) -> uuid.UUID:
    """Build an engagement whose findings are ``(title, severity, status, include_in_report)``."""
    with session_factory() as db:
        # A unique name per call: `scribble_clients.name` is UNIQUE, and a test that builds two
        # engagements (the mixed-vs-all-new comparison) would otherwise collide with itself.
        client = Client(name=f"Acme Co {uuid.uuid4().hex[:8]}")
        db.add(client)
        db.flush()
        eng = Engagement(name="Q3 Assessment", client_id=client.id, company_name="Acme Corp")
        group = FindingGroup(engagement=eng, name="External", order_index=0)
        for i, (title, severity, status, included) in enumerate(findings):
            EngagementFinding(
                engagement=eng,
                group=group,
                title=title,
                severity=severity,
                status=status,
                include_in_report=included,
                order_index=i,
                content_json={"description": _block(f"Detail for {title}.")},
            )
        db.add(eng)
        db.commit()
        return eng.id


def _titles(ctx) -> set[str]:
    out = set()
    for group in ctx.groups:
        for f in group.findings:
            out.add(f.title)
            out.update(c.title for c in f.children)
    return out


# ── E1: the defect itself ────────────────────────────────────────────────────────────────────────


def test_a_remediated_or_false_positive_finding_does_not_inflate_the_risk_rating(session_factory):
    """E1. One live High, one Critical marked ``fixed``, one Critical marked ``false_positive``.

    Before lotek#618 this reported overall=critical / total=3 — the client's headline risk driven by a
    finding that was remediated and one that was never real.
    """
    eng_id = _engagement(
        session_factory,
        [
            ("SMB signing not required", Severity.high, FindingStatus.new, True),
            ("Domain admin over SMB", Severity.critical, FindingStatus.fixed, True),
            ("Phantom RCE", Severity.critical, FindingStatus.false_positive, True),
        ],
    )
    with session_factory() as db:
        ctx = build_report_context(db.get(Engagement, eng_id))

    assert ctx.rollup is not None
    assert ctx.rollup.overall == "high", "a fixed/false-positive Critical must not set the overall risk"
    assert ctx.rollup.total == 1, "only live findings are counted"
    assert ctx.rollup.counts["critical"] == 0
    assert ctx.rollup.disposition_counts == {"live": 1, "remediated": 1, "accepted": 0, "excluded": 1}


def test_the_narrative_counts_only_live_findings(session_factory):
    """E1. ``_build_narrative`` reads the rollup, so the executive summary inherited the same lie."""
    eng_id = _engagement(
        session_factory,
        [
            ("SMB signing not required", Severity.high, FindingStatus.new, True),
            ("Domain admin over SMB", Severity.critical, FindingStatus.fixed, True),
            ("Phantom RCE", Severity.critical, FindingStatus.false_positive, True),
        ],
    )
    with session_factory() as db:
        ctx = build_report_context(db.get(Engagement, eng_id))

    assert "1 finding" in ctx.narrative
    assert "3 findings" not in ctx.narrative


# ── E2: every surface agrees ─────────────────────────────────────────────────────────────────────


@pytest.mark.parametrize("status", list(FindingStatus))
def test_disposition_and_label_are_derived_in_one_place(status):
    """E2. The predicate and the label map are the single source every surface reads."""
    disposition, label, _rendered, _counted = EXPECTED[status]
    assert report_disposition(status) == disposition
    assert finding_status_label(status) == label


@pytest.mark.parametrize("status", list(FindingStatus))
def test_every_status_agrees_across_context_and_rollup(status, session_factory):
    """E2. One finding per status: does it reach the context, is it counted, what label does it carry?

    A finding that renders but is uncounted (``fixed``/``accepted_risk``) is the case a second copy of
    the predicate gets wrong — the card says Remediated while the banner still counts it.
    """
    disposition, label, rendered, counted = EXPECTED[status]
    eng_id = _engagement(
        session_factory, [("Weak TLS configuration", Severity.medium, status, True)]
    )
    with session_factory() as db:
        ctx = build_report_context(db.get(Engagement, eng_id))

    present = "Weak TLS configuration" in _titles(ctx)
    assert present is rendered, f"{status}: expected rendered={rendered}"
    assert ctx.rollup is not None
    assert ctx.rollup.total == (1 if counted else 0), f"{status}: expected counted={counted}"

    if rendered:
        finding = ctx.groups[0].findings[0]
        assert finding.status == status.value
        assert finding.disposition == disposition
        assert finding.status_label == label


@pytest.mark.parametrize("status", list(FindingStatus))
def test_include_in_report_remains_an_independent_veto(status, session_factory):
    """E2. ``include_in_report=False`` removes a finding whatever its status — the operator's explicit
    veto is not replaced by the derived disposition, it is ANDed with it."""
    eng_id = _engagement(
        session_factory, [("Weak TLS configuration", Severity.medium, status, False)]
    )
    with session_factory() as db:
        ctx = build_report_context(db.get(Engagement, eng_id))

    assert "Weak TLS configuration" not in _titles(ctx)
    assert ctx.rollup is not None
    assert ctx.rollup.total == 0


def test_html_shows_a_status_column_and_badges_only_when_some_finding_is_not_new(session_factory):
    """E2 + E3. The index gains a Status column and the cards gain badges — but only when there is
    something to say. An all-``new`` engagement must render exactly as it did before this feature."""
    mixed = _engagement(
        session_factory,
        [
            ("SMB signing not required", Severity.high, FindingStatus.new, True),
            ("Domain admin over SMB", Severity.critical, FindingStatus.fixed, True),
        ],
    )
    with session_factory() as db:
        html = render_report_html(build_report_context(db.get(Engagement, mixed)))

    assert "<th>Status</th>" in html
    assert "Remediated" in html

    plain = _engagement(
        session_factory, [("SMB signing not required", Severity.high, FindingStatus.new, True)]
    )
    with session_factory() as db:
        html_plain = render_report_html(build_report_context(db.get(Engagement, plain)))

    assert "<th>Status</th>" not in html_plain
    # MARKUP, not the class name: the report is one self-contained document, so its stylesheet is
    # always inlined whole and `.status-badge`'s *rules* are present either way. What must be absent
    # for an all-`new` engagement is any emitted badge or status cell — no chip, no column, and (see
    # the DOCX test) no status line in the .docx.
    assert '<span class="status-badge' not in html_plain
    assert 'class="ix-status' not in html_plain


def test_docx_carries_the_same_labels_as_the_html(session_factory):
    """E2. The DOCX is a peer deliverable, not an afterthought: same predicate, same labels."""
    eng_id = _engagement(
        session_factory,
        [
            ("SMB signing not required", Severity.high, FindingStatus.needs_retest, True),
            ("Domain admin over SMB", Severity.critical, FindingStatus.accepted_risk, True),
            ("Phantom RCE", Severity.critical, FindingStatus.false_positive, True),
        ],
    )
    with session_factory() as db:
        payload = render_report_docx(build_report_context(db.get(Engagement, eng_id)))

    text = _all_text(docx.Document(io.BytesIO(payload)))
    assert "Awaiting retest" in text
    assert "Risk accepted" in text
    assert "Phantom RCE" not in text, "an excluded finding must not reach the DOCX either"


# ── E5: a label must never assert conduct nobody recorded ────────────────────────────────────────


def test_no_label_asserts_a_verification_or_a_client_decision(session_factory):
    """E5. The rendered deliverable may say a finding is Remediated (an operator's recorded status);
    it may not say it was *verified*, or that a *client decided* to accept a risk — nothing in the
    data records either. Same rule as ``tests/test_report_standing_prose.py``."""
    for label in (finding_status_label(s) for s in FindingStatus):
        lowered = label.lower()
        assert "verif" not in lowered
        assert "client" not in lowered

    eng_id = _engagement(
        session_factory,
        [
            ("Domain admin over SMB", Severity.critical, FindingStatus.fixed, True),
            ("Legacy TLS accepted", Severity.low, FindingStatus.accepted_risk, True),
        ],
    )
    with session_factory() as db:
        html = render_report_html(build_report_context(db.get(Engagement, eng_id)))

    assert "Fixed (verified)" not in html
    assert "client decision" not in html.lower()
