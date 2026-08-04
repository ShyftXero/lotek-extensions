"""WS8: DOCX report renderer tests.

Mirrors ``tests/test_report_html.py``'s fixture shape (same two-group, varied-severity, one-excluded,
one-{{TARGET_HOST}} engagement) so the two deliverables are exercised the same way, then asserts the
rendered ``.docx`` (reopened with ``python-docx``) honors the frozen contracts: group + finding board
order, severity text, excluded-finding absence, variable resolution, evidence-image embedding, and
that malformed/edge-case content never crashes the renderer.
"""

from __future__ import annotations

import io

import docx
import pytest
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

from scribble.content import schema
from scribble.content.render_docx import html_to_richtext
from scribble.enums import ArtifactKind, ArtifactPlacement, Severity
from scribble.models import Artifact, Client, Engagement, EngagementFinding, FindingGroup
from scribble.reporting import build_report_context
from scribble.reporting.render_docx import render_report_docx

# A 1x1 black PNG — enough for python-docx's image-header sniffing to accept it as real image bytes.
_PNG_BYTES = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
    "53de0000000c4944415478da6360606000000005000166ff0f0e0000000049454e44ae426082"
)


def _block(text: str) -> dict:
    return schema.doc_from_text(text)


def _target_host_block() -> dict:
    return {
        "type": schema.DOC,
        "content": [
            {
                "type": schema.PARAGRAPH,
                "content": [
                    {"type": schema.TEXT, "text": "Affected host: "},
                    {"type": schema.VARIABLE, "attrs": {"key": "TARGET_HOST"}},
                ],
            }
        ],
    }


def _rich_block() -> dict:
    """A block exercising headings/lists/marks/links so the docx walker gets real coverage."""
    return {
        "type": schema.DOC,
        "content": [
            {
                "type": schema.HEADING,
                "attrs": {"level": 2},
                "content": [{"type": schema.TEXT, "text": "Impact"}],
            },
            {
                "type": schema.PARAGRAPH,
                "content": [
                    {"type": schema.TEXT, "text": "bold", "marks": [{"type": "bold"}]},
                    {"type": schema.TEXT, "text": " and "},
                    {
                        "type": schema.TEXT,
                        "text": "linked",
                        "marks": [{"type": "link", "attrs": {"href": "https://example.test"}}],
                    },
                ],
            },
            {
                "type": schema.BULLET_LIST,
                "content": [
                    {
                        "type": schema.LIST_ITEM,
                        "content": [
                            {
                                "type": schema.PARAGRAPH,
                                "content": [{"type": schema.TEXT, "text": "Item one"}],
                            }
                        ],
                    },
                    {
                        "type": schema.LIST_ITEM,
                        "content": [
                            {
                                "type": schema.PARAGRAPH,
                                "content": [{"type": schema.TEXT, "text": "Item two"}],
                            }
                        ],
                    },
                ],
            },
        ],
    }


def _build_engagement(session_factory) -> int:
    with session_factory() as db:
        client = Client(name="Acme Co")
        db.add(client)
        db.flush()
        eng = Engagement(name="Q3 Combined Assessment", client_id=client.id, company_name="Acme Corp")
        internal = FindingGroup(engagement=eng, name="Internal", order_index=0)
        external = FindingGroup(engagement=eng, name="External", order_index=1)

        EngagementFinding(
            engagement=eng,
            group=internal,
            title="Weak SMB Signing",
            severity=Severity.low,
            order_index=0,
            content_json={"description": _block("SMB signing is not required on several hosts.")},
        )
        EngagementFinding(
            engagement=eng,
            group=internal,
            title="Domain Admin Compromise",
            severity=Severity.critical,
            order_index=1,
            content_json={
                "description": _block("Full domain compromise was achieved via Kerberoasting."),
                "remediation": _rich_block(),
            },
        )
        EngagementFinding(
            engagement=eng,
            group=external,
            title="Reflected XSS",
            severity=Severity.medium,
            order_index=0,
            target_host="app.acme.test",
            target_port="443",
            content_json={"description": _target_host_block()},
        )
        hidden = EngagementFinding(
            engagement=eng,
            group=external,
            title="Excluded Finding",
            severity=Severity.high,
            order_index=1,
            content_json={"description": _block("Should never render.")},
        )
        hidden.include_in_report = False

        db.add(eng)
        db.commit()
        return eng.id


def _all_text(doc: docx.Document) -> str:
    """Every ``<w:t>`` text run in the document body, in document order — including inside tables,
    which ``Document.paragraphs`` alone would miss."""
    return "".join(t.text or "" for t in doc.element.body.iter(qn("w:t")))


def _cell_fill(cell) -> str | None:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn("w:shd"))
    return shd.get(qn("w:fill")) if shd is not None else None


def _para_style(doc: docx.Document, text: str) -> str:
    """The style *name* of the first top-level paragraph whose text equals ``text``. Raises if none —
    the assertion should see a real paragraph, not silently pass on absence."""
    for p in doc.paragraphs:
        if p.text == text:
            return p.style.name
    raise AssertionError(f"no top-level paragraph with text {text!r}")


def _para_index(doc: docx.Document, text: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if p.text == text:
            return i
    raise AssertionError(f"no top-level paragraph with text {text!r}")


def test_render_report_docx_contract(session_factory):
    eng_id = _build_engagement(session_factory)
    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx)

    assert isinstance(payload, bytes)
    assert len(payload) > 0

    # --- opens cleanly as a valid docx -------------------------------------------------------
    doc = docx.Document(io.BytesIO(payload))
    text = _all_text(doc)

    # --- group order: Internal (order_index=0) before External (order_index=1) -------------
    idx_internal = text.index("Internal")
    idx_external = text.index("External")
    assert idx_internal < idx_external

    # --- finding order within a group: worst-first (auto_severity default) -----------------
    idx_crit = text.index("Domain Admin Compromise")
    idx_low = text.index("Weak SMB Signing")
    assert idx_crit < idx_low

    # --- excluded finding never renders ------------------------------------------------------
    assert "Excluded Finding" not in text

    # --- severity labels present --------------------------------------------------------------
    assert "Critical" in text
    assert "Medium" in text
    assert "Low" in text

    # --- {{TARGET_HOST}} resolved, no raw template markup survives --------------------------
    assert "app.acme.test" in text
    assert "{{" not in text
    assert "}}" not in text
    assert "{%" not in text
    assert "%}" not in text

    # --- rich content (heading/list/link) made it through the HTML->docx walker -------------
    assert "Impact" in text
    assert "Item one" in text
    assert "Item two" in text
    assert "linked" in text

    # --- per-severity cell coloring (Jinja `{% cellbg %}` conditional, PLAN.md §9) -----------
    finding_tables = doc.tables[1:]  # tables[0] is the executive-summary counts table
    assert len(finding_tables) == 3  # Domain Admin (crit), Weak SMB (low), Reflected XSS (medium)
    assert _cell_fill(finding_tables[0].rows[0].cells[0]) == "B91C1C"  # critical
    assert _cell_fill(finding_tables[1].rows[0].cells[0]) == "CA8A04"  # low
    assert _cell_fill(finding_tables[2].rows[0].cells[0]) == "EA580C"  # medium


def test_render_report_docx_applies_real_paragraph_styles(session_factory):
    """Regression guard (review C1/C2): headings and list items in a finding body must carry their
    real Word styles, not silently fall back to Normal — ``<w:pStyle>`` resolves by *styleId*, not
    the spaced display name. Also asserts no stray empty paragraph splits a list item from its text.
    """
    eng_id = _build_engagement(session_factory)
    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx)

    doc = docx.Document(io.BytesIO(payload))

    # Group + finding titles use the template's heading styles.
    assert _para_style(doc, "Internal") == "Heading 1"  # group
    assert _para_style(doc, "Domain Admin Compromise") == "Heading 2"  # finding title

    # Content-block labels render as Heading 3.
    assert _para_style(doc, "Remediation") == "Heading 3"

    # Rich content from the HTML->docx walker keeps its styles (C1).
    assert _para_style(doc, "Impact") == "Heading 2"  # <h2> inside the remediation block
    assert _para_style(doc, "Item one") == "List Bullet"
    assert _para_style(doc, "Item two") == "List Bullet"

    # No empty paragraph between the two list items (C2: the list style must sit on the text-bearing
    # paragraph, not on an empty one).
    i_one, i_two = _para_index(doc, "Item one"), _para_index(doc, "Item two")
    assert i_two == i_one + 1, "unexpected paragraph(s) between the two list items"
    assert doc.paragraphs[i_one].text and doc.paragraphs[i_two].text  # neither is empty


def test_render_report_docx_renders_nested_children_compactly(session_factory):
    """D1 (docx mirror of the HTML renderer's compact per-host list): a parent finding's write-up
    renders once, in its own table row; its children render as an "Affected Hosts" list appended
    INLINE into that same finding's body -- they never get their own top-level finding table."""
    with session_factory() as db:
        eng = Engagement(name="Nested Docx Report", company_name="Acme")
        g = FindingGroup(engagement=eng, name="Internal", order_index=0)
        parent = EngagementFinding(
            engagement=eng,
            group=g,
            title="Kerberoastable Account",
            severity=Severity.high,
            order_index=0,
            content_json={"description": _block("Kerberoastable accounts were identified.")},
        )
        db.add(eng)
        db.flush()

        child_a = EngagementFinding(
            engagement=eng,
            group=g,
            title="Kerberoastable Account",
            severity=Severity.high,
            order_index=1,
            target_host="dc01.acme.test",
            # Every child promoted under the same vuln-DB template shares that template's
            # ``content_json`` verbatim with its parent (see ``scribble.promote.promote_job``) -- these
            # per-child blocks are deliberately DIFFERENT text here to prove the renderer does NOT read
            # them for the per-host row; the real per-host evidence comes from ``variables`` instead
            # (``EngagementFinding.variables``, filled by promote from the host's own facts).
            content_json={"description": _block("Should not get its own table.")},
            variables={"AFFECTED": "svc_sql"},
        )
        child_b = EngagementFinding(
            engagement=eng,
            group=g,
            title="Kerberoastable Account",
            severity=Severity.high,
            order_index=2,
            target_host="dc02.acme.test",
            content_json={"description": _block("Nor should this one.")},
            variables={"AFFECTED": "svc_web"},
        )
        child_a.parent_id = parent.id
        child_b.parent_id = parent.id
        db.add_all([child_a, child_b])
        db.commit()
        eng_id = eng.id

    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx)

    doc = docx.Document(io.BytesIO(payload))
    text = _all_text(doc)

    # Only ONE finding table -- the parent's. Children never get their own (their own facts-derived
    # evidence line still appears, but INLINE in the "Affected Hosts" list inside that same one table).
    finding_tables = doc.tables[1:]
    assert len(finding_tables) == 1

    assert "Kerberoastable Account" in text
    assert "Kerberoastable accounts were identified." in text

    assert "Affected Hosts (2)" in text
    assert "dc01.acme.test" in text
    assert "dc02.acme.test" in text
    assert "svc_sql" in text  # child_a's own facts-derived evidence line, not a content-block excerpt
    assert "svc_web" in text
    # Neither child's own content block is rendered anywhere -- confirming the evidence line comes
    # from ``variables``, not from a copy of (parent or child) descriptive text.
    assert "Should not get its own table." not in text
    assert "Nor should this one." not in text


def test_render_report_docx_includes_narrative(session_factory):
    """D2: the generated executive-summary narrative is threaded into the docx scalars and renders
    as real text in the executive-summary paragraph the template's ``build_default_docx.py`` adds."""
    eng_id = _build_engagement(session_factory)
    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx)

    doc = docx.Document(io.BytesIO(payload))
    text = _all_text(doc)
    assert ctx.narrative != ""
    assert ctx.narrative in text


def test_render_report_docx_empty_engagement(session_factory):
    with session_factory() as db:
        eng = Engagement(name="Clean Sweep", company_name="Acme")
        db.add(eng)
        db.commit()
        eng_id = eng.id

    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx)

    doc = docx.Document(io.BytesIO(payload))
    text = _all_text(doc)
    assert "No findings recorded for this engagement." in text
    assert "{{" not in text and "{%" not in text


def test_render_report_docx_no_content_block_degrades_gracefully(session_factory):
    with session_factory() as db:
        eng = Engagement(name="Bare Finding Co", company_name="Acme")
        group = FindingGroup(engagement=eng, name="Web App", order_index=0)
        EngagementFinding(
            engagement=eng,
            group=group,
            title="Untitled Content Finding",
            severity=Severity.info,
            order_index=0,
            content_json={},
        )
        db.add(eng)
        db.commit()
        eng_id = eng.id

    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx)  # must not raise

    doc = docx.Document(io.BytesIO(payload))
    text = _all_text(doc)
    assert "Untitled Content Finding" in text
    assert "No content." in text


def test_evidence_image_embeds_as_inline_shape(session_factory):
    with session_factory() as db:
        client = Client(name="Acme")
        db.add(client)
        db.flush()
        eng = Engagement(name="Assets Engagement", client_id=client.id, company_name="Acme")
        group = FindingGroup(engagement=eng, name="Web App", order_index=0)
        finding = EngagementFinding(
            engagement=eng,
            group=group,
            title="Stored XSS",
            severity=Severity.high,
            order_index=0,
            content_json={"description": _block("See evidence below.")},
        )
        db.add(eng)
        db.flush()
        artifact = Artifact(
            engagement=eng,
            finding=finding,
            kind=ArtifactKind.screenshot,
            placement=ArtifactPlacement.attached,
            filename="poc.png",
            content_type="image/png",
            storage_path="poc.png",
            caption="Proof of concept",
            order_index=0,
        )
        db.add(artifact)
        db.commit()
        eng_id = eng.id

    fake_files = {"poc.png": _PNG_BYTES}

    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        embedded_payload = render_report_docx(ctx, artifact_bytes=fake_files.get)
        not_embedded_payload = render_report_docx(ctx)  # artifact_bytes=None -> graceful skip

    embedded_doc = docx.Document(io.BytesIO(embedded_payload))
    assert len(embedded_doc.inline_shapes) == 1
    assert "Proof of concept" in _all_text(embedded_doc)

    not_embedded_doc = docx.Document(io.BytesIO(not_embedded_payload))
    assert len(not_embedded_doc.inline_shapes) == 0
    not_embedded_text = _all_text(not_embedded_doc)
    assert "not embedded" in not_embedded_text
    assert "poc.png" in not_embedded_text


def test_oversized_evidence_image_degrades_to_caption(session_factory, monkeypatch):
    """Review W2: an artifact whose bytes exceed the size ceiling must not be embedded (falls back to
    caption-only) so a runaway file can't bloat the render's memory footprint."""
    import scribble.reporting.render_docx as rdx

    with session_factory() as db:
        eng = Engagement(name="Big Assets", company_name="Acme")
        group = FindingGroup(engagement=eng, name="Web App", order_index=0)
        finding = EngagementFinding(
            engagement=eng,
            group=group,
            title="Huge Screenshot Finding",
            severity=Severity.medium,
            order_index=0,
            content_json={"description": _block("Big evidence.")},
        )
        db.add(eng)
        db.flush()
        db.add(
            Artifact(
                engagement=eng,
                finding=finding,
                kind=ArtifactKind.screenshot,
                placement=ArtifactPlacement.attached,
                filename="huge.png",
                content_type="image/png",
                storage_path="huge.png",
                caption="Oversized proof",
                order_index=0,
            )
        )
        db.commit()
        eng_id = eng.id

    # Drop the ceiling well below the payload so we exercise the cap without allocating 25 MiB.
    monkeypatch.setattr(rdx, "_MAX_EVIDENCE_BYTES", 8)
    oversized = {"huge.png": _PNG_BYTES}  # > 8 bytes

    with session_factory() as db:
        engagement = db.get(Engagement, eng_id)
        ctx = build_report_context(engagement)
        payload = render_report_docx(ctx, artifact_bytes=oversized.get)

    doc = docx.Document(io.BytesIO(payload))
    assert len(doc.inline_shapes) == 0  # oversized image was NOT embedded
    text = _all_text(doc)
    assert "not embedded" in text
    assert "huge.png" in text


# --- content/render_docx.py unit-level robustness (malformed HTML, never crash) -----------------


@pytest.fixture
def tpl(tmp_path):
    """A minimal real DocxTemplate with ``current_rendering_part`` set, matching how
    ``reporting/render_docx.py`` prepares one before converting content blocks."""
    doc = docx.Document()
    doc.add_paragraph("{{r body}}")
    path = tmp_path / "mini.docx"
    doc.save(str(path))
    t = DocxTemplate(str(path))
    t.init_docx()
    t.current_rendering_part = t.docx.part
    return t


def test_html_to_richtext_handles_malformed_html(tpl):
    malformed = "<p>Unclosed paragraph <strong>bold <em>nested</p><ul><li>loose li outside ul</div>"
    rt = html_to_richtext(malformed, tpl=tpl)  # must not raise
    assert rt.xml  # produced *something*


def test_html_to_richtext_unknown_tags_degrade_to_text(tpl):
    html = "<p>Known text</p><marquee>unknown tag content</marquee><script>alert(1)</script>"
    rt = html_to_richtext(html, tpl=tpl)
    assert "Known text" in rt.xml
    assert "unknown tag content" in rt.xml


def test_html_to_richtext_missing_image_degrades_to_placeholder(tpl):
    html = '<p>Before</p><img src="/nope.png" alt="missing"/><p>After</p>'
    rt = html_to_richtext(html, tpl=tpl, image_resolver=lambda _src: None)
    assert "Before" in rt.xml
    assert "After" in rt.xml
    assert "image" in rt.xml.lower()


def test_html_to_richtext_empty_input(tpl):
    rt = html_to_richtext("", tpl=tpl)
    assert rt.xml == ""
    rt_none = html_to_richtext(None, tpl=tpl)
    assert rt_none.xml == ""
